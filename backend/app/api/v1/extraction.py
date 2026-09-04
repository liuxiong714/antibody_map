import csv
import io
import logging
import re
import uuid
from collections import Counter
from datetime import datetime, timezone
from typing import Any
from urllib.parse import quote

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from pydantic import BaseModel, model_validator
from sqlalchemy import case, func, select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user, get_db, require_admin
from app.core.audit import log_audit
from app.core.term_normalizer import CHINA_PROVINCE_NAMES
from app.core.traceability_html import (
    datapoint_dict_to_trace,
    generate_traceability_html,
)
from app.models.data_point import DataPoint
from app.models.extraction_history import ExtractionHistory
from app.models.literature import Literature
from app.models.titer_table import TiterTable
from app.models.user import User
from app.schemas.common import ApiResponse, PagedResponse
from app.services.extraction_service import (
    compute_data_point_conflicts,
    get_extraction_history,
    get_extraction_results,
    get_extraction_status,
    review_data_points,
    trigger_extraction,
)
from app.services.literature_service import LOCAL_STORAGE_DIR
from app.tasks.celery_app import celery_app
from app.tasks.quality_task import score_data_point_task

router = APIRouter()
logger = logging.getLogger("uvicorn")

# 文件扩展名列表，用于从标题中去除
_EXTENSIONS_PATTERN = re.compile(r"\.(pdf|caj|doc|docx|txt|epub|pptx|xlsx|ps|wps|md)$", re.IGNORECASE)
# 年份前缀，如 "2025 ", "2024_", "2025-", "2025." 等
_YEAR_PREFIX_PATTERN = re.compile(r"^(19\d{2}|20\d{2})\s*[ _\-\.,;:]\s*")


def clean_literature_title(title: str) -> str:
    """清洗文献标题：去除文件后缀和年份前缀"""
    if not title:
        return title
    t = title.strip()
    # 1. 去除文件扩展名
    # 只去除末尾的扩展名，确保不是真正的标题中的点
    t = _EXTENSIONS_PATTERN.sub("", t).strip()
    # 2. 去除年份前缀
    # 重复应用以处理 "2025 2024 Title" 这种多重前缀
    while True:
        new_t = _YEAR_PREFIX_PATTERN.sub("", t).strip()
        if new_t == t:
            break
        t = new_t
    # 3. 去除末尾多余的空格/标点
    t = t.strip(" ._-,;:")
    return t if t else title  # 如果清洗后为空则保留原标题


# ── 请求体模型 ──────────────────────────────────────────

_DP_DATA_TYPES = {"seroprevalence", "gmc"}
_DP_CONFIDENCES = {"high", "medium", "low"}


def _validate_datapoint_consistency(
    data_type: str | None,
    value: float | None,
    ci_lower: float | None,
    ci_upper: float | None,
    sample_size: int | None,
    age_min: float | None,
    age_max: float | None,
    collection_year: int | None,
    confidence: str | None,
    source_char_start: int | None,
    source_char_end: int | None,
) -> None:
    """P0-5：数据点字段级一致性校验，杜绝阳性率 150%、age_max<age_min、年份 9999 等脏数据入库。

    仅对显式给出的字段校验；字段为 None 视为"不修改/未知"，跳过。非法抛 ValueError，
    由 FastAPI 统一转为 422 响应。
    """
    if data_type is not None:
        if data_type not in _DP_DATA_TYPES:
            raise ValueError(
                f"data_type 必须是 {sorted(_DP_DATA_TYPES)} 之一，当前为 {data_type!r}"
            )
        if value is not None:
            if data_type == "seroprevalence" and not (0.0 <= value <= 100.0):
                raise ValueError(f"血清阳性率 value 必须在 [0,100] 区间，当前为 {value!r}")
            if data_type == "gmc" and value < 0:
                raise ValueError(f"GMC value 不能为负，当前为 {value!r}")
    if ci_lower is not None and ci_upper is not None and ci_upper < ci_lower:
        raise ValueError(f"CI 下限不能大于上限（lower={ci_lower} > upper={ci_upper}）")
    if sample_size is not None and sample_size < 0:
        raise ValueError(f"sample_size 不能为负，当前为 {sample_size!r}")
    if age_min is not None and age_min < 0:
        raise ValueError(f"age_min 不能为负，当前为 {age_min!r}")
    if age_max is not None and age_max < 0:
        raise ValueError(f"age_max 不能为负，当前为 {age_max!r}")
    if age_min is not None and age_max is not None and age_max < age_min:
        raise ValueError(f"age_max({age_max!r}) 不能小于 age_min({age_min!r})")
    if collection_year is not None:
        y = int(collection_year)
        if not (1900 <= y <= 2100):
            raise ValueError(f"collection_year 必须在 [1900,2100] 区间，当前为 {collection_year!r}")
    if confidence is not None and confidence not in _DP_CONFIDENCES:
        raise ValueError(f"confidence 必须是 {sorted(_DP_CONFIDENCES)} 之一，当前为 {confidence!r}")
    if source_char_start is not None and source_char_start < 0:
        raise ValueError(f"source_char_start 不能为负，当前为 {source_char_start!r}")
    if source_char_end is not None and source_char_end < 0:
        raise ValueError(f"source_char_end 不能为负，当前为 {source_char_end!r}")
    if source_char_start is not None and source_char_end is not None and source_char_end < source_char_start:
        raise ValueError(
            f"source_char_end({source_char_end!r}) 不能小于 source_char_start({source_char_start!r})"
        )

class DataPointReviewItem(BaseModel):
    id: str
    review_status: str | None = None  # "approved" | "rejected" | None (仅编辑时不审核)
    review_note: str | None = None
    review_comment: str | None = None  # 审核意见
    # 以下为可编辑的数据字段
    disease: str | None = None
    province: str | None = None
    city: str | None = None
    data_type: str | None = None
    value: float | None = None
    unit: str | None = None
    sample_size: int | None = None
    population: str | None = None
    age_min: float | None = None
    age_max: float | None = None
    collection_year: int | None = None
    confidence: str | None = None
    method: str | None = None
    assay: str | None = None
    source_page: int | None = None
    source_context: str | None = None
    # P0 新增：精确字符级溯源
    source_char_start: int | None = None
    source_char_end: int | None = None
    is_grounded: bool | None = None

    @model_validator(mode="after")
    def _check_consistency(self):
        _validate_datapoint_consistency(
            self.data_type, self.value, None, None,
            self.sample_size, self.age_min, self.age_max,
            self.collection_year, self.confidence,
            self.source_char_start, self.source_char_end,
        )
        return self


class UpdateDataPointsRequest(BaseModel):
    data_points: list[DataPointReviewItem]


class BatchReviewRequest(BaseModel):
    ids: list[str]
    note: str | None = None  # 兼容历史字段
    comment: str | None = None  # 审核意见（驳回时必填）


class ExtractionRequest(BaseModel):
    model: str | None = None
    model_config_id: str | None = None
    api_key: str | None = None
    base_url: str | None = None
    clear_existing_data: bool = False
    # 是否使用 Redis 提取结果缓存；False 时强制重新提取（跳过 LLM 缓存）
    use_cache: bool = True


class BatchExtractionRequest(BaseModel):
    """批量重新提取请求"""
    literature_ids: list[str]
    model: str | None = None
    model_config_id: str | None = None
    api_key: str | None = None
    base_url: str | None = None
    clear_existing_data: bool = False
    # 是否使用 Redis 提取结果缓存；False 时强制重新提取
    use_cache: bool = True


class CreateDataPointRequest(BaseModel):
    """手动新增数据点"""
    disease: str | None = None
    province: str | None = None
    city: str | None = None
    data_type: str | None = None
    value: float | None = None
    unit: str | None = None
    sample_size: int | None = None
    population: str | None = None
    age_min: float | None = None
    age_max: float | None = None
    collection_year: int | None = None
    confidence: str | None = "medium"
    method: str | None = None
    assay: str | None = None
    source_page: int | None = None
    source_context: str | None = None
    # P0 新增：精确字符级溯源
    source_char_start: int | None = None
    source_char_end: int | None = None
    is_grounded: bool = False

    @model_validator(mode="after")
    def _check_consistency(self):
        _validate_datapoint_consistency(
            self.data_type, self.value, None, None,
            self.sample_size, self.age_min, self.age_max,
            self.collection_year, self.confidence,
            self.source_char_start, self.source_char_end,
        )
        return self


# ── 提取相关路由 ────────────────────────────────────────

@router.post("/literatures/{literature_id}/extraction", response_model=ApiResponse, summary="触发AI提取", description="触发单篇文献的AI数据提取任务，可指定模型、API Key和Base URL，支持清空已有数据后重新提取")
async def start_extraction(
    literature_id: uuid.UUID,
    req: ExtractionRequest = None,
    db: AsyncSession = Depends(get_db),
):
    """触发文献 AI 数据提取任务"""
    try:
        model = req.model if req else None
        model_config_id = req.model_config_id if req else None
        api_key = req.api_key if req else None
        base_url = req.base_url if req else None
        clear_existing = req.clear_existing_data if req else False
        use_cache = req.use_cache if req else True
        result = await trigger_extraction(db, literature_id, model, api_key, base_url, model_config_id, clear_existing, use_cache)
        return ApiResponse(message="提取任务已提交", data=result)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@router.post("/literatures/extraction/batch", response_model=ApiResponse, summary="批量触发AI提取", description="批量触发多篇文献的AI数据提取任务：有 PDF 走全文提取，无 PDF 但有摘要的题录文献用摘要提取；仅无 PDF 且无摘要、或正在提取中的文献会被跳过")
async def start_batch_extraction(
    req: BatchExtractionRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量触发文献 AI 数据提取任务"""
    if not req.literature_ids:
        raise HTTPException(status_code=400, detail="请选择至少一个文献")

    # 普通用户每次批量提交受数量上限约束；管理员不受限
    MAX_BATCH = 100
    if not current_user.is_admin and len(req.literature_ids) > MAX_BATCH:
        raise HTTPException(
            status_code=400,
            detail=f"普通用户单次最多提交 {MAX_BATCH} 篇，当前选择了 {len(req.literature_ids)} 篇",
        )

    submitted: list[dict] = []
    skipped: list[dict] = []
    errors: list[dict] = []

    for lit_id_str in req.literature_ids:
        try:
            lit_id = uuid.UUID(lit_id_str)
            # 检查文献存在且状态非 processing 或 queued
            result = await db.execute(
                select(Literature).where(Literature.id == lit_id)
            )
            literature = result.scalar_one_or_none()
            if not literature:
                errors.append({"id": lit_id_str, "reason": "文献不存在"})
                continue
            # 有 PDF 走全文提取，无 PDF 但有摘要时用摘要提取；两者皆无才跳过
            if not literature.file_path and not literature.abstract:
                skipped.append({"id": lit_id_str, "title": literature.title, "reason": "既无 PDF 也无摘要，无法提取"})
                continue
            if literature.extraction_status in ("processing", "queued"):
                skipped.append({"id": lit_id_str, "title": literature.title, "reason": f"当前状态 {literature.extraction_status}，跳过"})
                continue

            # 触发提取
            await trigger_extraction(
                db, lit_id,
                model=req.model,
                api_key=req.api_key,
                base_url=req.base_url,
                model_config_id=req.model_config_id,
                clear_existing_data=req.clear_existing_data,
                use_cache=req.use_cache,
            )
            submitted.append({
                "id": str(lit_id),
                "title": literature.title,
            })
        except ValueError as e:
            errors.append({"id": lit_id_str, "reason": str(e)})
        except Exception as e:
            logger.error(f"[BatchExtract] 提交文献 {lit_id_str} 失败: {e}", exc_info=True)
            errors.append({"id": lit_id_str, "reason": str(e)})

    logger.info(
        f"[BatchExtract] 批量提取提交完成: 提交{len(submitted)}篇, "
        f"跳过{len(skipped)}篇, 失败{len(errors)}篇"
    )

    return ApiResponse(
        message=f"批量提取提交完成：成功 {len(submitted)} 篇，跳过 {len(skipped)} 篇，失败 {len(errors)} 篇",
        data={
            "submitted": submitted,
            "skipped": skipped,
            "errors": errors,
            "submitted_count": len(submitted),
            "skipped_count": len(skipped),
            "error_count": len(errors),
        },
    )


@router.get("/literatures/{literature_id}/extraction/status", response_model=ApiResponse, summary="查询提取状态", description="查询指定文献的AI数据提取任务当前状态（pending/processing/done/failed等）")
async def check_status(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """查询提取任务状态"""
    try:
        result = await get_extraction_status(db, literature_id)
        return ApiResponse(data=result)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@router.get("/literatures/{literature_id}/extraction", response_model=ApiResponse, summary="获取提取结果", description="获取指定文献的AI提取数据点列表及提取状态")
async def get_results(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """获取提取的数据点列表"""
    try:
        data_points = await get_extraction_results(db, literature_id)
        status = await get_extraction_status(db, literature_id)
        return ApiResponse(
            data={
                "literature_id": str(literature_id),
                "status": status["status"],
                "data_points": data_points,
            }
        )
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@router.get("/literatures/{literature_id}/extraction/export", summary="导出数据点CSV", description="将指定文献的所有数据点导出为CSV文件，便于查看和分析")
async def export_data_points(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """导出文献的数据点为 CSV"""
    data_points = await get_extraction_results(db, literature_id)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "疾病", "省份", "城市", "数据类型", "数值", "单位", "样本量",
        "年龄下限", "年龄上限", "采集年份", "置信度", "审核状态", "来源页码",
        "原文依据", "溯源区间", "是否已匹配原文",
    ])
    for dp in data_points:
        interval = ""
        if dp.get("source_char_start") is not None and dp.get("source_char_end") is not None:
            interval = f"[{dp['source_char_start']}, {dp['source_char_end']})"
        writer.writerow([
            dp.get("disease", ""),
            dp.get("province", ""),
            dp.get("city", ""),
            dp.get("data_type", ""),
            dp.get("value"),
            dp.get("unit", ""),
            dp.get("sample_size"),
            dp.get("age_min"),
            dp.get("age_max"),
            dp.get("collection_year"),
            dp.get("confidence", ""),
            dp.get("review_status", ""),
            dp.get("source_page", ""),
            (dp.get("source_context") or "").replace("\n", " "),
            interval,
            "是" if dp.get("is_grounded") else "否",
        ])

    return Response(
        content=output.getvalue().encode("utf-8-sig"),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''data_points_{literature_id}.csv"},
    )


@router.get("/literatures/{literature_id}/extraction/export-word", summary="导出数据点Word报告", description="将指定文献的数据点导出为Word报告（.docx格式），包含文献元信息、数据摘要、数据点明细表格")
async def export_data_points_word(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """导出文献的数据点为 Word 报告（.docx）"""
    # 1. 文献信息
    lit_result = await db.execute(
        select(Literature).where(Literature.id == literature_id)
    )
    literature = lit_result.scalar_one_or_none()
    if not literature:
        raise HTTPException(status_code=404, detail="文献不存在")

    data_points = await get_extraction_results(db, literature_id)

    doc = Document()

    # ── 标题 ──
    title_text = literature.title or f"文献 {literature_id}"
    doc.add_heading(title_text, level=0)

    # ── 文献元信息 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    meta_info = []
    if literature.authors:
        meta_info.append(f"作者：{literature.authors}")
    if literature.journal:
        meta_info.append(f"期刊：{literature.journal}")
    if literature.pub_year:
        meta_info.append(f"发表年份：{literature.pub_year}")
    if literature.extraction_status:
        status_map = {"processing": "进行中", "done": "已完成", "done_no_data": "完成（无数据）", "failed": "失败", "pending": "待处理"}
        meta_info.append(f"提取状态：{status_map.get(literature.extraction_status, literature.extraction_status)}")
    p.add_run(" | ".join(meta_info)).font.size = Pt(10)

    # ── 分隔线 ──
    doc.add_paragraph("─" * 50)

    # ── 统计摘要 ──
    doc.add_heading("数据摘要", level=1)
    total = len(data_points)
    if total > 0:
        by_type = Counter(dp.get("data_type", "") for dp in data_points)
        by_status = Counter(dp.get("review_status", "") for dp in data_points)
        provinces = {dp.get("province", "") for dp in data_points if dp.get("province")}
        diseases = {dp.get("disease", "") for dp in data_points if dp.get("disease")}

        summary = doc.add_table(rows=7, cols=2, style="Light List Accent 1")
        summary.cell(0, 0).text = "统计项"
        summary.cell(0, 1).text = "数值"
        summary.cell(1, 0).text = "数据点总数"
        summary.cell(1, 1).text = str(total)
        summary.cell(2, 0).text = "血清阳性率"
        summary.cell(2, 1).text = str(by_type.get("seroprevalence", 0))
        summary.cell(3, 0).text = "GMC"
        summary.cell(3, 1).text = str(by_type.get("gmc", 0))
        summary.cell(4, 0).text = "覆盖省份"
        summary.cell(4, 1).text = f"{len(provinces)} 个（{', '.join(sorted(provinces))}）"
        summary.cell(5, 0).text = "涉及疾病"
        summary.cell(5, 1).text = f"{len(diseases)} 种（{', '.join(sorted(diseases))}）"
        summary.cell(6, 0).text = "审核状态"
        summary.cell(6, 1).text = f"已通过 {by_status.get('approved', 0)} | 待审核 {by_status.get('pending', 0)} | 已驳回 {by_status.get('rejected', 0)}"
    else:
        doc.add_paragraph("暂无数据点")

    # ── 数据点明细 ──
    doc.add_heading("数据点明细", level=1)
    if data_points:
        table = doc.add_table(rows=1 + len(data_points), cols=9, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        headers = ["疾病", "省份", "数据类型", "数值", "样本量", "年龄段", "采集年份", "置信度", "审核状态"]
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # 数据行
        for idx, dp in enumerate(data_points):
            age_range = ""
            if dp.get("age_min") is not None or dp.get("age_max") is not None:
                age_min = dp.get("age_min") or ""
                age_max = dp.get("age_max") or ""
                age_range = f"{age_min}-{age_max}" if age_min and age_max else f"{age_min or ''}~{age_max or ''}"
            values = [
                dp.get("disease", ""),
                dp.get("province", ""),
                "阳性率" if dp.get("data_type") == "seroprevalence" else "GMC" if dp.get("data_type") == "gmc" else dp.get("data_type", ""),
                f"{dp.get('value', '')}{dp.get('unit', '')}",
                str(dp.get("sample_size", "") or ""),
                age_range,
                str(dp.get("collection_year", "") or ""),
                dp.get("confidence", ""),
                {"approved": "通过", "rejected": "驳回", "pending": "待审"}.get(dp.get("review_status", ""), dp.get("review_status", "")),
            ]
            for col, val in enumerate(values):
                cell = table.cell(idx + 1, col)
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
    else:
        doc.add_paragraph("暂无数据点")

    # ── 生成时间 ──
    doc.add_paragraph("─" * 50)
    ts = doc.add_paragraph()
    ts_run = ts.add_run(f"报告生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
    ts_run.font.size = Pt(9)
    ts_run.font.color.rgb = RGBColor(128, 128, 128)

    # ── 输出 ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c for c in title_text if c not in r'\/:*?"<>|').strip() or str(literature_id)
    filename = quote(f"{safe_title}_数据报告.docx")

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


@router.get("/literatures/{literature_id}/extraction/traceability-html", summary="导出溯源HTML报告", description="导出自包含溯源HTML文件，高亮原文并附带侧边栏数据点列表，可离线打开查看")
async def export_traceability_html(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """P1-2：导出自包含溯源 HTML（高亮原文 + 侧边栏数据点列表，可离线打开）"""
    # 1. 文献标题
    lit_result = await db.execute(
        select(Literature.title).where(Literature.id == literature_id)
    )
    title_row = lit_result.scalar_one_or_none()
    if not title_row:
        raise HTTPException(status_code=404, detail="文献不存在")
    title = title_row or str(literature_id)

    # 2. 数据点列表
    data_points = await get_extraction_results(db, literature_id)
    if not data_points:
        raise HTTPException(status_code=404, detail="该文献暂无提取数据点，无法生成溯源 HTML")

    # 3. 读取缓存的 clean_text（溯源文本）
    text_path = LOCAL_STORAGE_DIR / f"{literature_id}.txt"
    if not text_path.exists():
        raise HTTPException(
            status_code=404,
            detail="溯源文本未缓存，请重新提取该文献后再导出 HTML",
        )
    try:
        full_text = text_path.read_text(encoding="utf-8")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取溯源文本失败: {e}") from e

    # 4. 转换数据点为 TracePoint 并生成 HTML
    traces = [datapoint_dict_to_trace(dpo) for dpo in data_points]
    html_content = generate_traceability_html(
        title=title,
        full_text=full_text,
        data_points=traces,
    )

    # 5. 构建安全的下载文件名
    safe_title = "".join(c for c in title if c not in r'\/:*?"<>|').strip() or str(literature_id)
    filename = quote(f"{safe_title}_溯源报告.html")

    logger.info(
        f"[P1-2] 导出溯源 HTML: literature_id={literature_id}, "
        f"dp_count={len(traces)}, text_len={len(full_text)}"
    )

    return Response(
        content=html_content.encode("utf-8"),
        media_type="text/html; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


@router.post("/literatures/{literature_id}/extraction/data-points", response_model=ApiResponse, summary="手动添加数据点", description="为指定文献手动添加一个数据点，可设置所有数据字段和溯源信息")
async def create_data_point(
    literature_id: uuid.UUID,
    req: CreateDataPointRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """手动新增数据点"""
    # 验证文献存在
    result = await db.execute(
        select(Literature).where(Literature.id == literature_id)
    )
    literature = result.scalar_one_or_none()
    if not literature:
        raise HTTPException(status_code=404, detail="文献不存在")

    # 创建数据点
    dp = DataPoint(
        literature_id=literature_id,
        disease=req.disease,
        province=req.province,
        city=req.city,
        data_type=req.data_type,
        value=req.value,
        unit=req.unit,
        sample_size=req.sample_size,
        population=req.population,
        age_min=req.age_min,
        age_max=req.age_max,
        collection_year=req.collection_year,
        confidence=req.confidence or "medium",
        method=req.method,
        assay=req.assay,
        source_page=req.source_page,
        source_context=req.source_context,
        # P0 新增：精确字符级溯源
        source_char_start=req.source_char_start,
        source_char_end=req.source_char_end,
        is_grounded=bool(req.is_grounded),
        review_status="pending",
    )
    db.add(dp)
    await db.flush()

    # 更新文献提取状态和计数
    newly_done = False
    if literature.extraction_status in (None, "", "failed", "pending"):
        literature.extraction_status = "done"
        newly_done = True
    literature.extracted_count = (literature.extracted_count or 0) + 1
    literature.updated_at = datetime.now(timezone.utc)

    # 手动补录数据点把状态置为 done 时同步写历史，避免「有结果无历史」缺口
    if newly_done:
        db.add(
            ExtractionHistory(
                literature_id=literature_id,
                model=literature.llm_model_used,
                status="success",
                data_point_count=literature.extracted_count or 1,
                error_message="手动补录数据点：文献提取状态由 failed/pending 置为 done",
                extracted_at=literature.updated_at,
            )
        )

    await db.commit()

    # 4.2：记录数据点新增审计（独立会话自行提交，失败不影响主流程）
    await log_audit(
        db,
        "data_point_create",
        user_id=str(current_user.id),
        username=current_user.username,
        target=f"literature/{literature_id}",
        detail={"literature_id": str(literature_id), "disease": req.disease, "value": req.value},
        entity_type="data_point",
        entity_id=str(dp.id),
        new_value={
            "disease": dp.disease,
            "province": dp.province,
            "city": dp.city,
            "data_type": dp.data_type,
            "value": dp.value,
            "unit": dp.unit,
            "sample_size": dp.sample_size,
            "population": dp.population,
            "confidence": dp.confidence,
            "method": dp.method,
            "assay": dp.assay,
            "collection_year": dp.collection_year,
        },
    )

    return ApiResponse(
        message="数据点已添加",
        data={"id": str(dp.id)},
    )


# ── 审核相关路由 ────────────────────────────────────────

@router.put("/literatures/{literature_id}/extraction", response_model=ApiResponse, summary="更新数据点", description="批量更新数据点，可编辑任意字段和审核状态，支持编辑数据字段和审核状态")
async def update_data_points(
    literature_id: uuid.UUID,
    req: UpdateDataPointsRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """更新数据点（可编辑任意字段 + 审核状态）"""
    updated = []
    editable_fields = [
        "disease", "province", "city", "data_type", "value", "unit",
        "sample_size", "population", "age_min", "age_max", "collection_year",
        "confidence", "method", "assay", "source_page", "source_context",
        # P0 新增：精确字符级溯源
        "source_char_start", "source_char_end", "is_grounded",
    ]
    # 4.2：可审计字段（可编辑字段 + 审核状态 + 审核意见）
    audit_fields = [*editable_fields, "review_status", "review_comment"]

    # 4.2：预取将被更新的数据点旧值，用于生成变更 diff（审计）
    ids_to_update = [uuid.UUID(item.id) for item in req.data_points]
    old_rows = []
    if ids_to_update:
        old_result = await db.execute(
            select(DataPoint).where(
                DataPoint.id.in_(ids_to_update),
                DataPoint.literature_id == literature_id,
            )
        )
        old_rows = old_result.scalars().all()
    old_by_id = {str(d.id): d for d in old_rows}

    now = datetime.now(timezone.utc)
    # 4.2：待写审计条目（提交后再落库，独立会话）
    audit_entries: list[dict[str, Any]] = []

    for item in req.data_points:
        # 构建要更新的字段
        values: dict[str, Any] = {}
        dp_id = uuid.UUID(item.id)
        old = old_by_id.get(str(item.id))

        # 审核状态（写入审核人/审核时间）
        if item.review_status:
            if item.review_status not in ("approved", "rejected"):
                raise HTTPException(status_code=400, detail=f"无效的审核状态: {item.review_status}")
            values["review_status"] = item.review_status
            values["reviewer_id"] = current_user.id
            values["reviewed_at"] = now

        # 审核意见（仅当显式传入时更新，支持清空）
        if "review_comment" in item.model_dump(exclude_unset=True):
            values["review_comment"] = item.review_comment

        # 可编辑的数据字段（仅更新显式传入的字段，None 值表示清空）
        explicit = item.model_dump(exclude_unset=True, exclude={"id", "review_status", "review_note"})
        for field in editable_fields:
            if field in explicit:
                values[field] = explicit[field]

        if not values:
            continue

        stmt = (
            update(DataPoint)
            .where(DataPoint.id == dp_id)
            .where(DataPoint.literature_id == literature_id)
            .values(**values)
        )
        await db.execute(stmt)
        updated.append(item.id)

        # 4.2：生成本次变更 diff（仅记录实际变化的字段）
        if old is not None:
            old_diff: dict[str, Any] = {}
            new_diff: dict[str, Any] = {}
            for field in audit_fields:
                if field in values:
                    ov = getattr(old, field, None)
                    nv = values[field]
                    # 兼容 UUID（reviewer_id 等）
                    if isinstance(ov, uuid.UUID):
                        ov = str(ov)
                    if isinstance(nv, uuid.UUID):
                        nv = str(nv)
                    if ov != nv:
                        old_diff[field] = ov
                        new_diff[field] = nv
            audit_entries.append(
                {
                    "action": "data_point_review" if item.review_status else "data_point_update",
                    "entity_id": str(dp_id),
                    "old_diff": old_diff,
                    "new_diff": new_diff,
                }
            )

    # 如果有审核状态变更，同步 literature.approved_count（修复审核状态显示不正确的问题）
    has_review_change = any(item.review_status for item in req.data_points)
    if has_review_change:
        await _sync_approved_count(db, literature_id)

    await db.commit()

    # 4.2：数据点变更审计（独立会话，失败降级不影响主流程）
    for entry in audit_entries:
        await log_audit(
            db,
            entry["action"],
            user_id=str(current_user.id),
            username=current_user.username,
            target=f"literature/{literature_id}",
            entity_type="data_point",
            entity_id=entry["entity_id"],
            old_value=entry["old_diff"],
            new_value=entry["new_diff"],
        )

    # 审核通过后异步质量打分（幂等，全文可用后精打覆盖）
    approved_ids = [
        item.id for item in req.data_points if item.review_status == "approved"
    ]
    for dp_id in approved_ids:
        score_data_point_task.delay(dp_id)

    return ApiResponse(message="数据点已更新", data={"updated": updated})


@router.post("/literatures/{literature_id}/extraction/confirm", response_model=ApiResponse, summary="批量审核通过", description="批量将多个数据点审核通过，写入审核人/审核时间，同步更新文献的已通过计数")
async def batch_confirm(
    literature_id: uuid.UUID,
    req: BatchReviewRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量审核通过"""
    comment = req.comment if req.comment is not None else req.note
    result = await review_data_points(
        db, literature_id, req.ids, "approved", comment, current_user.id
    )

    await _sync_approved_count(db, literature_id)
    await db.commit()

    # 4.2：批量审核通过审计
    for dp_id in req.ids:
        await log_audit(
            db,
            "data_point_review",
            user_id=str(current_user.id),
            username=current_user.username,
            target=f"literature/{literature_id}",
            detail={"review_status": "approved", "review_comment": comment},
            entity_type="data_point",
            entity_id=str(dp_id),
            new_value={"review_status": "approved", "review_comment": comment},
        )

    # 审核通过后异步质量打分（幂等，全文可用后精打覆盖）
    for dp_id in req.ids:
        score_data_point_task.delay(dp_id)

    return ApiResponse(message=f"已批量通过 {result} 个数据点")


@router.post("/literatures/{literature_id}/extraction/dispute", response_model=ApiResponse, summary="批量驳回", description="批量驳回多个数据点，驳回必须填写意见，写入审核人/审核时间，同步更新文献的已通过计数")
async def batch_dispute(
    literature_id: uuid.UUID,
    req: BatchReviewRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量驳回"""
    comment = req.comment if req.comment is not None else req.note
    if not comment or not str(comment).strip():
        raise HTTPException(status_code=400, detail="驳回必须填写审核意见")

    result = await review_data_points(
        db, literature_id, req.ids, "rejected", comment, current_user.id
    )

    await _sync_approved_count(db, literature_id)
    await db.commit()

    # 4.2：批量驳回审计
    for dp_id in req.ids:
        await log_audit(
            db,
            "data_point_review",
            user_id=str(current_user.id),
            username=current_user.username,
            target=f"literature/{literature_id}",
            detail={"review_status": "rejected", "review_comment": comment},
            entity_type="data_point",
            entity_id=str(dp_id),
            new_value={"review_status": "rejected", "review_comment": comment},
        )

    return ApiResponse(message=f"已批量驳回 {result} 个数据点", data={"note": comment})


async def _sync_approved_count(db: AsyncSession, literature_id: uuid.UUID):
    """同步文献表中 approved_count"""
    count_result = await db.execute(
        select(func.count(DataPoint.id))
        .where(DataPoint.literature_id == literature_id)
        .where(DataPoint.review_status == "approved")
    )
    approved = count_result.scalar() or 0

    await db.execute(
        update(Literature)
        .where(Literature.id == literature_id)
        .values(approved_count=approved, updated_at=datetime.now(timezone.utc))
    )


@router.post("/literatures/{literature_id}/sync-metadata", response_model=ApiResponse, summary="同步文献元数据", description="从数据点聚合同步文献的pub_year和province元数据，并清洗标题（去除文件后缀和年份前缀）")
async def sync_literature_metadata(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """从数据点聚合同步文献的 pub_year 和 province 元数据，并清洗标题"""
    lit_result = await db.execute(
        select(Literature).where(Literature.id == literature_id)
    )
    literature = lit_result.scalar_one_or_none()
    if not literature:
        raise HTTPException(status_code=404, detail="文献不存在")

    # 清洗标题
    title_updated = False
    original_title = literature.title
    cleaned_title = clean_literature_title(original_title)
    if cleaned_title != original_title:
        literature.title = cleaned_title
        title_updated = True
        logger.info(f"[MetadataSync] 文献 {literature_id} 清洗标题: '{original_title}' -> '{cleaned_title}'")

    # 查询该文献的所有数据点
    dp_result = await db.execute(
        select(DataPoint).where(DataPoint.literature_id == literature_id)
    )
    all_data_points = dp_result.scalars().all()

    if not all_data_points:
        if title_updated:
            literature.updated_at = datetime.now(timezone.utc)
            await db.commit()
            await db.refresh(literature)
        return ApiResponse(message="元数据同步完成（标题已清洗）" if title_updated else "无需同步：该文献暂无数据点", data={
            "title_updated": title_updated,
            "pub_year_updated": False,
            "province_updated": False,
            "data_point_count": 0,
        })

    pub_year_updated = False
    province_updated = False

    # 聚合 pub_year
    if not literature.pub_year:
        years = []
        for dp in all_data_points:
            y = dp.collection_year
            if y:
                years.append(y)
        if years:
            literature.pub_year = max(set(years), key=years.count)
            pub_year_updated = True
            logger.info(f"[MetadataSync] 文献 {literature_id} 聚合更新 pub_year={literature.pub_year}")

    # 聚合 province
    if not literature.province:
        provinces = [
            dp.province for dp in all_data_points
            if dp.province and dp.province in CHINA_PROVINCE_NAMES
        ]
        if provinces:
            literature.province = max(set(provinces), key=provinces.count)
            province_updated = True
            logger.info(
                f"[MetadataSync] 文献 {literature_id} 聚合更新 province={literature.province} "
                f"(覆盖{len(set(provinces))}省)"
            )

    if pub_year_updated or province_updated or title_updated:
        literature.updated_at = datetime.now(timezone.utc)
        await db.commit()
        await db.refresh(literature)

    return ApiResponse(
        message="元数据同步完成",
        data={
            "id": str(literature.id),
            "title": literature.title,
            "title_updated": title_updated,
            "pub_year": literature.pub_year,
            "province": literature.province,
            "pub_year_updated": pub_year_updated,
            "province_updated": province_updated,
            "data_point_count": len(all_data_points),
        },
    )


@router.post("/literatures/sync-metadata-batch", response_model=ApiResponse, summary="批量同步文献元数据", description="批量同步所有提取完成的文献元数据，包括清洗标题和从数据点聚合pub_year和province")
async def sync_metadata_batch(db: AsyncSession = Depends(get_db)):
    """批量同步所有提取完成的文献元数据，包括：
    - 清洗标题：去除文件后缀（.pdf/.caj等）和年份前缀（2025、2024等）
    - 从数据点聚合 pub_year 和 province
    """
    # 查询所有 extraction_status='done' 的文献
    result = await db.execute(
        select(Literature).where(Literature.extraction_status == "done")
    )
    literatures = result.scalars().all()

    if not literatures:
        return ApiResponse(message="无需同步：没有已提取完成的文献", data={
            "total": 0, "synced": 0, "skipped": 0, "details": [],
        })

    # 查询这些文献的所有数据点
    lit_ids = [lit.id for lit in literatures]
    dp_result = await db.execute(
        select(DataPoint).where(DataPoint.literature_id.in_(lit_ids))
    )
    all_dps = dp_result.scalars().all()

    # 按 literature_id 分组
    dp_map: dict[uuid.UUID, list[DataPoint]] = {}
    for dp in all_dps:
        dp_map.setdefault(dp.literature_id, []).append(dp)

    synced_count = 0
    skipped_count = 0
    details = []

    for literature in literatures:
        dps = dp_map.get(literature.id, [])
        if not dps:
            skipped_count += 1
            continue

        pub_year_updated = False
        province_updated = False
        title_updated = False

        # 清洗标题：去除文件后缀和年份前缀
        original_title = literature.title
        cleaned_title = clean_literature_title(original_title)
        if cleaned_title != original_title:
            literature.title = cleaned_title
            title_updated = True
            logger.info(
                f"[MetadataSync-Batch] 文献 {literature.id} 清洗标题: "
                f"'{original_title}' -> '{cleaned_title}'"
            )

        if not literature.pub_year:
            years = [dp.collection_year for dp in dps if dp.collection_year]
            if years:
                literature.pub_year = max(set(years), key=years.count)
                pub_year_updated = True

        if not literature.province:
            provinces = [
                dp.province for dp in dps
                if dp.province and dp.province in CHINA_PROVINCE_NAMES
            ]
            if provinces:
                literature.province = max(set(provinces), key=provinces.count)
                province_updated = True

        if pub_year_updated or province_updated or title_updated:
            literature.updated_at = datetime.now(timezone.utc)
            synced_count += 1
            details.append({
                "id": str(literature.id),
                "title": literature.title,
                "title_updated": title_updated,
                "pub_year": literature.pub_year,
                "province": literature.province,
                "pub_year_updated": pub_year_updated,
                "province_updated": province_updated,
            })
        else:
            skipped_count += 1

    if synced_count > 0:
        await db.commit()

    logger.info(
        f"[MetadataSync-Batch] 批量同步完成: 共{len(literatures)}篇, "
        f"同步{synced_count}篇, 跳过{skipped_count}篇"
    )

    return ApiResponse(
        message=f"批量同步完成：{synced_count}篇已更新，{skipped_count}篇无需更新",
        data={
            "total": len(literatures),
            "synced": synced_count,
            "skipped": skipped_count,
            "details": details,
        },
    )


# ── 提取状态修复与手动停止 ──────────────────────────────────

@router.post("/literatures/{literature_id}/extraction/stop", response_model=ApiResponse, summary="停止提取", description="手动停止/重置文献提取状态，将extraction_status=processing的文献强制重置为failed，用于处理因服务器重启等导致的永久提取中状态")
async def stop_extraction(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """手动停止/重置文献提取状态。

    将 extraction_status='processing' 的文献强制重置为 'failed'，
    用于处理服务器重启或任务异常退出导致的永久「提取中」状态。
    """
    result = await db.execute(
        select(Literature).where(Literature.id == literature_id)
    )
    literature = result.scalar_one_or_none()
    if not literature:
        raise HTTPException(status_code=404, detail="文献不存在")

    if literature.extraction_status != "processing":
        return ApiResponse(
            message=f"当前状态为 {literature.extraction_status}，无需停止",
            data={"literature_id": str(literature_id), "status": literature.extraction_status},
        )

    literature.extraction_status = "failed"
    literature.updated_at = datetime.now(timezone.utc)
    # 手动停止提取时同步写失败历史，避免「有结果无历史」缺口
    db.add(
        ExtractionHistory(
            literature_id=literature_id,
            model=literature.llm_model_used,
            status="failed",
            data_point_count=0,
            error_message="手动停止：文献提取状态由 processing 重置为 failed",
            extracted_at=literature.updated_at,
        )
    )
    await db.commit()

    logger.warning(f"文献 {literature_id} 提取已被手动停止，状态重置为 failed")
    return ApiResponse(
        message="提取已停止，状态重置为失败",
        data={"literature_id": str(literature_id), "status": "failed"},
    )


@router.get("/literatures/{literature_id}/extraction/history", response_model=ApiResponse, summary="获取提取历史", description="获取指定文献的历次AI提取历史记录")
async def get_history(
    literature_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """获取文献的历次 AI 提取历史"""
    try:
        history = await get_extraction_history(db, literature_id)
        return ApiResponse(data=history)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@router.post("/literatures/extraction/reset-stuck", response_model=ApiResponse, summary="批量重置卡住的提取（管理员）", description="管理员专用：批量重置所有卡在processing或queued状态的文献为failed，并强制终止运行中的Celery提取任务，清空队列，用于服务器重启后恢复状态")
async def reset_stuck_extractions(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """批量重置所有卡在 'processing' 或 'queued' 状态的文献为 'failed'，
    同时强制终止运行中的 Celery 提取任务并清空队列。

    典型场景：服务器重启后，之前的异步提取任务已丢失但状态未更新。
    """
    # 1. 强制终止所有运行中的 Celery 提取任务
    purged_count = 0
    revoked_count = 0
    try:
        # 查看当前活跃的任务（timeout=5 避免默认1秒超时导致查不到任务）
        inspect = celery_app.control.inspect(timeout=5.0)
        active_tasks = inspect.active() or {}
        for _worker, tasks in active_tasks.items():
            for task in tasks:
                task_id = task.get("id")
                if task_id:
                    celery_app.control.revoke(task_id, terminate=True)
                    revoked_count += 1
                    logger.warning(f"[ResetStuck] 强制终止运行中任务: {task_id} ({task.get('name', 'unknown')})")

        # 清空队列中所有等待的任务
        purged_count = celery_app.control.purge()
        if purged_count > 0:
            logger.warning(f"[ResetStuck] 清空队列 {purged_count} 个等待中的任务")
    except Exception as e:
        logger.warning(f"[ResetStuck] Celery 控制操作失败（不影响数据库重置）: {e}")

    # 2. 重置数据库状态
    result = await db.execute(
        select(Literature).where(Literature.extraction_status.in_(["processing", "queued"]))
    )
    stuck_list = result.scalars().all()

    reset_ids = []
    for lit in stuck_list:
        lit.extraction_status = "failed"
        lit.updated_at = datetime.now(timezone.utc)
        reset_ids.append(str(lit.id))
        # 重置卡住状态时同步写失败历史，避免「有结果无历史」缺口
        db.add(
            ExtractionHistory(
                literature_id=lit.id,
                model=lit.llm_model_used,
                status="failed",
                data_point_count=0,
                error_message="管理员批量重置：提取状态 processing/queued 被重置为 failed",
                extracted_at=lit.updated_at,
            )
        )

    await db.commit()

    message_parts = []
    if reset_ids:
        message_parts.append(f"已重置 {len(reset_ids)} 篇文献状态为失败")
    if purged_count > 0:
        message_parts.append(f"已清空 {purged_count} 个排队任务")
    if revoked_count > 0:
        message_parts.append(f"已终止 {revoked_count} 个运行中任务")

    if not message_parts:
        message_parts.append("没有卡住的提取任务")

    final_message = "，".join(message_parts)

    logger.warning(
        f"[ResetStuck] {final_message} | "
        f"reset_ids={reset_ids}, purged={purged_count}, revoked={revoked_count}"
    )
    return ApiResponse(
        message=final_message,
        data={
            "reset_count": len(reset_ids),
            "literature_ids": reset_ids,
            "purged_count": purged_count,
            "revoked_count": revoked_count,
        },
    )


@router.post("/literatures/extraction/reset-my", response_model=ApiResponse, summary="终止我的提取（仅限自己提交的文献）", description="终止当前登录用户自己提交（owner_id 匹配）的卡在 processing/queued 状态的 AI 提取任务，重置为 failed，并尽力终止运行中/排队中属于本人的任务。不影响其他用户提交的提取任务。")
async def reset_my_extractions(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """终止仅属于当前用户（owner_id == current_user.id）的卡住提取任务。

    仅能影响自己提交（归属）的文献，不触碰其他用户的提取任务，也不做全局清队列。
    """
    # 1. 找出属于自己的卡住文献（processing / queued）
    result = await db.execute(
        select(Literature).where(
            Literature.owner_id == current_user.id,
            Literature.extraction_status.in_(["processing", "queued"]),
        )
    )
    own_stuck = result.scalars().all()
    own_ids = [lit.id for lit in own_stuck]
    own_id_strs = {str(lit_id) for lit_id in own_ids}

    revoked_count = 0
    if own_id_strs:
        # 2. 尽力终止运行中/排队中属于本人文献的 Celery 提取任务
        try:
            inspect = celery_app.control.inspect(timeout=5.0)
            active_tasks = inspect.active() or {}
            for _worker, tasks in active_tasks.items():
                for task in tasks:
                    task_name = task.get("name") or ""
                    if not task_name.endswith("process_literature"):
                        continue
                    args = task.get("args") or ()
                    kwargs = task.get("kwargs") or {}
                    lit_ref = ""
                    if args:
                        lit_ref = str(args[0])
                    elif kwargs.get("literature_id"):
                        lit_ref = str(kwargs["literature_id"])
                    if lit_ref in own_id_strs:
                        task_id = task.get("id")
                        if task_id:
                            celery_app.control.revoke(task_id, terminate=True)
                            revoked_count += 1
                            logger.warning(f"[ResetMy] 终止属于本人的运行中任务: {task_id} (文献 {lit_ref})")
        except Exception as e:
            logger.warning(f"[ResetMy] Celery 控制操作失败（不影响数据库重置）: {e}")

    # 3. 重置本人文献状态为 failed
    reset_ids = []
    for lit in own_stuck:
        lit.extraction_status = "failed"
        lit.updated_at = datetime.now(timezone.utc)
        reset_ids.append(str(lit.id))
        # 终止本人卡住提取时同步写失败历史，避免「有结果无历史」缺口
        db.add(
            ExtractionHistory(
                literature_id=lit.id,
                model=lit.llm_model_used,
                status="failed",
                data_point_count=0,
                error_message="终止(reset-my)：本人提交的 processing/queued 提取被重置为 failed",
                extracted_at=lit.updated_at,
            )
        )

    await db.commit()

    message_parts = []
    if reset_ids:
        message_parts.append(f"已重置 {len(reset_ids)} 篇属于您的文献状态为失败")
    if revoked_count > 0:
        message_parts.append(f"已终止 {revoked_count} 个属于您的运行中任务")
    if not message_parts:
        message_parts.append("没有属于您的卡住提取任务")

    final_message = "，".join(message_parts)
    logger.warning(
        f"[ResetMy] {final_message} | user={current_user.username}, reset_ids={reset_ids}, revoked={revoked_count}"
    )
    return ApiResponse(
        message=final_message,
        data={
            "reset_count": len(reset_ids),
            "literature_ids": reset_ids,
            "revoked_count": revoked_count,
            "purged_count": 0,
        },
    )


# ── 提取队列状态查询 ──────────────────────────────────

@router.get("/extractions/queue-status", response_model=ApiResponse, summary="查询提取队列状态", description="查询当前AI提取队列状态：待处理数、排队中数、提取中数等，以及排队中的文献列表")
async def get_extraction_queue_status(db: AsyncSession = Depends(get_db)):
    """查询AI提取队列状态，用于前端展示当前工作负载"""
    # 先重置卡死状态，与列表查询口径一致，避免两者数字不一致
    from app.services.literature.crud import reset_stale_extraction_status
    await reset_stale_extraction_status(db)

    # 各状态计数
    count_result = await db.execute(
        select(Literature.extraction_status, func.count(Literature.id))
        .where(Literature.deleted_at.is_(None))
        .group_by(Literature.extraction_status)
    )
    counts = {row[0]: row[1] for row in count_result.fetchall()}

    # 排队中的文献列表（queued）
    queued_result = await db.execute(
        select(Literature.id, Literature.title)
        .where(Literature.extraction_status == "queued")
        .where(Literature.deleted_at.is_(None))
        .order_by(Literature.updated_at)
        .limit(100)
    )
    queued_list = [
        {"id": str(row[0]), "title": row[1] or ""}
        for row in queued_result.fetchall()
    ]

    # 提取中的文献列表（processing）
    processing_result = await db.execute(
        select(Literature.id, Literature.title)
        .where(Literature.extraction_status == "processing")
        .where(Literature.deleted_at.is_(None))
        .order_by(Literature.updated_at)
        .limit(100)
    )
    processing_list = [
        {"id": str(row[0]), "title": row[1] or ""}
        for row in processing_result.fetchall()
    ]

    return ApiResponse(data={
        "pending_count": counts.get("pending", 0),
        "queued_count": counts.get("queued", 0),
        "processing_count": counts.get("processing", 0),
        "done_count": counts.get("done", 0),
        "done_no_data_count": counts.get("done_no_data", 0),
        "failed_count": counts.get("failed", 0),
        "total": sum(counts.values()),
        "queued_literatures": queued_list,
        "processing_literatures": processing_list,
    })


# ---------------------------------------------------------------------------
# F15：审核队列（跨文献待审/已驳回数据点，按置信度与质量分排序，低置信优先）
# ---------------------------------------------------------------------------
_CONFIDENCE_RANK = case(
    (DataPoint.confidence == "high", 3),
    (DataPoint.confidence == "medium", 2),
    (DataPoint.confidence == "low", 1),
    else_=2,
)


def _serialize_review_dp(dp: DataPoint, literature_title: str | None, conflicts: list | None = None) -> dict:
    """审核队列数据点序列化（与 get_extraction_results 字段对齐，附加文献标题）。"""
    return {
        "id": str(dp.id),
        "literature_id": str(dp.literature_id),
        "literature_title": literature_title or "",
        "disease": dp.disease,
        "region": dp.region,
        "province": dp.province,
        "city": dp.city,
        "data_type": dp.data_type,
        "value": float(dp.value) if dp.value else None,
        "unit": dp.unit,
        "ci_lower": float(dp.ci_lower) if dp.ci_lower else None,
        "ci_upper": float(dp.ci_upper) if dp.ci_upper else None,
        "sample_size": dp.sample_size,
        "method": dp.method,
        "assay": dp.assay,
        "population": dp.population,
        "age_min": dp.age_min,
        "age_max": dp.age_max,
        "collection_year": dp.collection_year,
        "source_page": dp.source_page,
        "source_context": dp.source_context,
        "is_grounded": bool(dp.is_grounded),
        "confidence": dp.confidence,
        "review_status": dp.review_status,
        "review_comment": dp.review_comment,
        "quality_score": dp.quality_score,
        "quality_grade": dp.quality_grade,
        "estimate_grade": dp.estimate_grade,
        # P1-6：同省同病同年已有已审核数据点冲突对比（审核页只读提示）
        "conflicts": conflicts or [],
        "created_at": dp.created_at.isoformat() if dp.created_at else None,
    }


@router.get("/extractions/review-queue", response_model=PagedResponse, summary="审核队列",
            description="跨文献列出待审核(pending)/已驳回(rejected)的数据点，按置信度（低→高）与质量分（低→高）排序，低置信低质量优先审核")
async def get_review_queue(
    review_status: str | None = Query(None, description="pending/rejected/all；默认 pending"),
    data_type: str | None = Query(None, description="seroprevalence/gmc"),
    disease: str | None = Query(None, description="疾病关键词模糊匹配"),
    literature_id: uuid.UUID | None = Query(None, description="限定文献"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    status = (review_status or "pending").strip().lower()
    if status not in ("pending", "rejected", "all"):
        raise HTTPException(status_code=400, detail="无效的审核状态（pending/rejected/all）")

    base = (
        select(DataPoint, Literature.title)
        .join(Literature, DataPoint.literature_id == Literature.id)
        .where(Literature.deleted_at.is_(None))
    )
    if status == "all":
        base = base.where(DataPoint.review_status.in_(["pending", "rejected"]))
    else:
        base = base.where(DataPoint.review_status == status)
    if data_type:
        base = base.where(DataPoint.data_type == data_type)
    if disease:
        base = base.where(DataPoint.disease.ilike(f"%{disease}%"))
    if literature_id:
        base = base.where(DataPoint.literature_id == literature_id)

    total = (
        await db.execute(
            select(func.count())
            .select_from(base.subquery())
        )
    ).scalar_one()

    rows = (
        await db.execute(
            base.order_by(
                _CONFIDENCE_RANK.asc(),
                DataPoint.quality_score.asc().nulls_last(),
                DataPoint.created_at.asc(),
            )
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
    ).all()

    # P1-6：同省同病同年已有已审核数据点冲突对比（审核页只读提示）
    conflict_map: dict = {}
    try:
        conflict_map = await compute_data_point_conflicts(db, [dp for dp, _ in rows])
    except Exception as e:  # 冲突检测失败不应影响审核队列加载
        logger.warning(f"P1-6 审核队列冲突检测失败（忽略）: {e}")
    items = [
        _serialize_review_dp(dp, title, conflict_map.get(str(dp.id), []))
        for dp, title in rows
    ]

    return PagedResponse(
        items=items, total=total, page=page, page_size=page_size,
        message=f"审核队列：{total} 条",
    )


# ---------------------------------------------------------------------------
# F16：滴度矩阵审核（titer_table 落库后的人工审核衔接）
# ---------------------------------------------------------------------------
@router.get("/titer-tables/review-queue", response_model=PagedResponse, summary="滴度矩阵审核队列",
            description="跨文献列出待审核/已驳回的滴度矩阵，按置信度（低→高）排序，供人工审核")
async def get_titer_review_queue(
    review_status: str | None = Query(None, description="pending/rejected/all；默认 pending"),
    assay_type: str | None = Query(None, description="hi/vnt/elisa"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    status = (review_status or "pending").strip().lower()
    if status not in ("pending", "rejected", "all"):
        raise HTTPException(status_code=400, detail="无效的审核状态（pending/rejected/all）")

    base = (
        select(TiterTable, Literature.title)
        .join(Literature, TiterTable.literature_id == Literature.id)
        .where(Literature.deleted_at.is_(None))
    )
    if status == "all":
        base = base.where(TiterTable.review_status.in_(["pending", "rejected"]))
    else:
        base = base.where(TiterTable.review_status == status)
    if assay_type:
        base = base.where(TiterTable.assay_type == assay_type)

    total = (
        await db.execute(select(func.count()).select_from(base.subquery()))
    ).scalar_one()

    titer_rank = case(
        (TiterTable.confidence == "high", 3),
        (TiterTable.confidence == "medium", 2),
        (TiterTable.confidence == "low", 1),
        else_=2,
    )
    rows = (
        await db.execute(
            base.order_by(titer_rank.asc(), TiterTable.created_at.asc())
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
    ).all()

    items = [
        {
            "id": str(tt.id),
            "literature_id": str(tt.literature_id),
            "literature_title": title or "",
            "assay_type": tt.assay_type,
            "ref_antisera": tt.ref_antisera,
            "antigens": tt.antigens,
            "titers": tt.titers,
            "unit": tt.unit,
            "quality_score": tt.quality_score,
            "source_page": tt.source_page,
            "source_context": tt.source_context,
            "confidence": tt.confidence,
            "review_status": tt.review_status,
            "review_comment": tt.review_comment,
            "created_at": tt.created_at.isoformat() if tt.created_at else None,
        }
        for tt, title in rows
    ]
    return PagedResponse(
        items=items, total=total, page=page, page_size=page_size,
        message=f"滴度矩阵审核队列：{total} 条",
    )


class TiterReviewRequest(BaseModel):
    review_status: str  # approved / rejected
    review_comment: str | None = None


@router.put("/titer-tables/{titer_table_id}/review", response_model=ApiResponse, summary="审核滴度矩阵",
            description="审核通过/驳回指定滴度矩阵，写入审核状态与审核意见")
async def review_titer_table(
    titer_table_id: uuid.UUID,
    req: TiterReviewRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if req.review_status not in ("approved", "rejected"):
        raise HTTPException(status_code=400, detail="无效的审核状态（approved/rejected）")
    if req.review_status == "rejected" and not (req.review_comment or "").strip():
        raise HTTPException(status_code=400, detail="驳回必须填写审核意见")

    tt = await db.get(TiterTable, titer_table_id)
    if tt is None:
        raise HTTPException(status_code=404, detail="滴度矩阵不存在")

    old_status = tt.review_status
    tt.review_status = req.review_status
    tt.review_comment = req.review_comment
    await db.commit()

    try:
        await log_audit(
            db, "titer_table_review",
            user_id=str(current_user.id),
            entity_type="titer_table", entity_id=str(tt.id),
            old_value={"review_status": old_status},
            new_value={"review_status": req.review_status, "review_comment": req.review_comment},
        )
    except Exception as e:
        logger.warning(f"滴度矩阵审核审计日志写入失败: {e}")

    return ApiResponse(
        message="审核通过" if req.review_status == "approved" else "已驳回",
        data={"id": str(tt.id), "review_status": tt.review_status},
    )
