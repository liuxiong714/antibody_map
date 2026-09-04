import base64
import contextlib
import hashlib
import json
import logging
from datetime import datetime, timezone
from uuid import UUID

from openai import AsyncOpenAI
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.methodology import build_methodology_note
from app.models.api_model_config import ApiModelConfig
from app.models.data_point import DataPoint
from app.models.report import Report
from app.models.report_template import ReportTemplate

logger = logging.getLogger("uvicorn")


def _data_snapshot_hash(rows) -> str | None:
    """对报告所依据的审核通过数据点生成稳定指纹。

    取每条记录的关键字段（id/literature_id/省份/疾病/年龄段/样本量/阳性值/数据类型/审核状态），
    排序后序列化为 JSON 再 SHA-256，得到对底层数据集敏感的固定 hash，用于可复现性校验。
    """
    if not rows:
        return None
    canonical = sorted(
        (
            str(getattr(r, "id", "")), str(getattr(r, "literature_id", "")),
            str(getattr(r, "province", "") or ""), str(getattr(r, "disease", "") or ""),
            getattr(r, "age_min", None), getattr(r, "age_max", None),
            getattr(r, "sample_size", None), getattr(r, "value", None),
            str(getattr(r, "data_type", "") or ""), str(getattr(r, "review_status", "") or ""),
        )
        for r in rows
    )
    payload = json.dumps(canonical, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()

DISEASE_NAMES = {
    "measles": "麻疹", "mumps": "腮腺炎", "rubella": "风疹",
    "pertussis": "百日咳", "diphtheria": "白喉", "tetanus": "破伤风",
    "hepatitis_b": "乙肝", "hepatitis_a": "甲肝", "polio": "脊灰",
    "influenza": "流感", "covid19": "新冠", "meningitis": "流脑",
    "varicella": "水痘", "hfmd": "手足口病", "rotavirus": "轮状病毒",
}

AGE_GROUPS = [
    ("<1岁", 0, 0),
    ("1-4岁", 1, 4),
    ("5-14岁", 5, 14),
    ("15-59岁", 15, 59),
    (">=60岁", 60, 200),
]

AGE_GROUPS_EN = [
    ("<1y", 0, 0),
    ("1-4y", 1, 4),
    ("5-14y", 5, 14),
    ("15-59y", 15, 59),
    (">=60y", 60, 200),
]

REPORT_PROMPT_ZH = """你是一位流行病学专家，请根据以下数据生成一份{title}的抗体水平分析报告。

数据概况：
- 数据来源：{literature_count} 篇文献
- 数据点数：{point_count} 个
- 覆盖省份：{province_count} 个
- 总样本量：{total_samples} 人

各省数据：
{province_table}

年份趋势：
{year_trend}

年龄分布：
{age_distribution}

请按以下结构输出报告（Markdown 格式）：
## 1. 总体概况
## 2. 地区分布
## 3. 时间趋势分析
## 4. 年龄分布特征
## 5. 免疫学参考意见

请基于数据给出专业的免疫学分析和建议，不要编造不存在的数据。"""

REPORT_PROMPT_EN = """You are an epidemiological expert. Please generate an antibody level analysis report based on the following data for {title_en}.

Data Overview:
- Data Sources: {literature_count} papers
- Data Points: {point_count}
- Provinces Covered: {province_count}
- Total Sample Size: {total_samples}

Province Data:
{province_table}

Yearly Trend:
{year_trend}

Age Distribution:
{age_distribution}

Please structure the report in Markdown format:
## 1. Overall Summary
## 2. Regional Distribution
## 3. Temporal Trend Analysis
## 4. Age Distribution Characteristics
## 5. Immunological Recommendations

Base your analysis strictly on the provided data. Do not fabricate any data."""

IMMUNE_BARRIER_PROMPT_ZH = """你是一位免疫学和流行病学专家。请根据以下审核通过的人群抗体水平数据，撰写一份{title}的免疫屏障评估报告。

数据概况：
- 数据来源：{literature_count} 篇文献
- 数据点数：{point_count} 个
- 覆盖省份：{province_count} 个
- 总样本量：{total_samples} 人

各省数据：
{province_table}

年份趋势：
{year_trend}

年龄分布：
{age_distribution}

请按以下结构输出报告（Markdown 格式）：
## 1. 免疫屏障总体评估
- 基于人群抗体阳性率与样本量综合判断整体免疫屏障水平
## 2. 地区免疫屏障差异分析
- 识别免疫屏障薄弱地区与高风险地区
- 分析地区间抗体水平差异及其公共卫生意义
## 3. 时间维度屏障变化趋势
- 结合年份趋势判断免疫屏障的时间演变规律
## 4. 年龄分段屏障特征
- 分析不同年龄段的抗体水平与免疫缺口
## 5. 免疫屏障缺口识别与干预建议
- 综合评估人群免疫保护状况，指出薄弱环节与高风险人群
- 给出加强免疫、监测与防控的具体建议

请基于数据给出专业的免疫屏障评估，不要编造不存在的数据。"""

IMMUNE_BARRIER_PROMPT_EN = """You are an expert in immunology and epidemiology. Based on the following reviewed population antibody titer data, write an {title_en} immune barrier assessment report.

Data Overview:
- Data Sources: {literature_count} papers
- Data Points: {point_count}
- Provinces Covered: {province_count}
- Total Sample Size: {total_samples}

Province Data:
{province_table}

Yearly Trend:
{year_trend}

Age Distribution:
{age_distribution}

Please structure the report in Markdown format:
## 1. Overall Immune Barrier Assessment
## 2. Regional Immune Barrier Differences
## 3. Temporal Trends of Barrier Changes
## 4. Age-stratified Barrier Characteristics
## 5. Barrier Gaps and Intervention Recommendations

Base your analysis strictly on the provided data. Do not fabricate any data."""

VACCINATION_STRATEGY_PROMPT_ZH = """你是一位流行病学和疫苗学专家。请根据以下任务信息和任务地点传染病流行情况，综合研判并制定参加任务人员的疫苗接种策略。

任务信息：
- 任务类型：{task_type}
- 任务时间：{task_time}
- 任务地点：{task_location}
- 人员人数：{personnel_count}人
- 人员性别分布：{personnel_gender}
- 人员年龄范围：{personnel_age}
- 人员疫苗接种史：{personnel_vaccination_history}

任务地点传染病流行情况：
{epidemic_data}

请按以下结构输出疫苗接种策略报告（Markdown 格式）：
## 1. 任务概况与传染病风险评估
- 结合任务类型、时间、地点分析可能面临的传染病风险
- 评估任务地点的传染病流行威胁
## 2. 任务地点传染病流行现状分析
- 基于现有流行病学数据分析当地主要传染病的流行水平和免疫屏障状况
- 识别高风险传染病
## 3. 人员免疫状态评估
- 结合人员年龄、性别、疫苗接种史评估群体免疫水平
- 识别免疫缺口
## 4. 推荐疫苗接种方案
- 按优先级排列推荐接种的疫苗种类
- 说明每项疫苗的推荐理由（结合当地流行情况和人员特点）
- 给出建议接种率目标
## 5. 接种时间安排建议
- 根据任务时间倒推接种窗口
- 多剂次疫苗的接种排程
## 6. 其他防护措施建议
- 除疫苗接种外的健康防护建议

请基于数据给出专业的疫苗接种策略建议，确保建议具有可操作性。不要编造不存在的数据。"""


def _calc_weighted_rate(rows: list[DataPoint]) -> tuple[float, int]:
    """计算加权阳性率"""
    sp_rows = [r for r in rows if r.data_type == "seroprevalence" and r.sample_size and r.value is not None]
    if not sp_rows:
        return 0.0, 0
    total_sample = sum(r.sample_size for r in sp_rows)
    weighted_sum = sum(r.value * r.sample_size for r in sp_rows)
    return round(weighted_sum / total_sample, 2), total_sample


async def _resolve_model_name(db: AsyncSession, model: str | None = None) -> str:
    """解析模型显示名称"""
    if not model:
        return settings.LLM_MODEL
    try:
        uid = UUID(model)
        result = await db.execute(select(ApiModelConfig).where(ApiModelConfig.id == uid))
        config = result.scalar_one_or_none()
        if config:
            return f"{config.name} ({config.model_name})"
    except (ValueError, AttributeError):
        pass
    return model


DEFAULT_TEMPLATES = [
    {
        "name": "抗体水平分析报告（默认）",
        "report_type": "antibody_analysis",
        "is_default": True,
        "desc": "含数据概况、时间/地区/年龄分析、总体概况与免疫学建议的标准抗体分析报告。",
        "sections": [
            {"title": "数据概览", "type": "kpi", "order": 0, "kpi": ["literature_count", "point_count", "province_count", "total_samples", "weighted_rate"]},
            {"title": "总体概况", "type": "text", "order": 1, "content_template": "请基于数据概况与分省/分年/分年龄摘要，撰写报告《总体概况》章节，给出整体抗体水平的判断与解读。"},
            {"title": "时间趋势分析", "type": "chart", "order": 2, "analysis": "trend"},
            {"title": "地区分布分析", "type": "chart", "order": 3, "analysis": "region"},
            {"title": "年龄分布特征", "type": "chart", "order": 4, "analysis": "age_curve"},
            {"title": "免疫学参考意见", "type": "text", "order": 5, "content_template": "请基于以上数据撰写《免疫学参考意见》章节，给出专业的免疫学分析与防控/接种建议，不要编造不存在的数据。"},
        ],
    },
    {
        "name": "疫苗接种策略报告（默认）",
        "report_type": "vaccination_strategy",
        "is_default": True,
        "desc": "含任务概况、疾病流行现状、年龄分布与接种方案建议的策略研判报告。",
        "sections": [
            {"title": "任务与数据概况", "type": "kpi", "order": 0, "kpi": ["literature_count", "point_count", "province_count", "weighted_rate"]},
            {"title": "传染病风险评估", "type": "text", "order": 1, "content_template": "请结合任务类型、时间、地点与当地传染病流行数据，撰写《传染病风险评估》章节，识别高风险传染病。"},
            {"title": "疾病流行现状分析", "type": "chart", "order": 2, "analysis": "disease"},
            {"title": "人群年龄分布特征", "type": "chart", "order": 3, "analysis": "age_curve"},
            {"title": "推荐疫苗接种方案", "type": "text", "order": 4, "content_template": "请基于当地流行情况与人员特点，撰写《推荐疫苗接种方案》章节，按优先级给出可操作性建议。"},
        ],
    },
    {
        "name": "免疫屏障评估报告（默认）",
        "report_type": "immune_barrier_assessment",
        "is_default": True,
        "desc": "含数据概览、免疫屏障总体/地区/时间/年龄评估与缺口干预建议的评估报告。",
        "sections": [
            {"title": "数据概览", "type": "kpi", "order": 0, "kpi": ["literature_count", "point_count", "province_count", "total_samples", "weighted_rate"]},
            {"title": "免疫屏障总体评估", "type": "text", "order": 1, "content_template": "请基于数据概况与分省/分年/分年龄摘要，撰写《免疫屏障总体评估》章节，综合判断人群整体免疫屏障水平。"},
            {"title": "地区免疫屏障差异分析", "type": "chart", "order": 2, "analysis": "region"},
            {"title": "时间维度屏障变化趋势", "type": "chart", "order": 3, "analysis": "trend"},
            {"title": "年龄分段屏障特征", "type": "chart", "order": 4, "analysis": "age_curve"},
            {"title": "免疫屏障缺口与干预建议", "type": "text", "order": 5, "content_template": "请基于以上数据撰写《免疫屏障缺口与干预建议》章节，识别免疫薄弱环节与高风险人群，并给出加强免疫、监测与防控建议，不要编造不存在的数据。"},
        ],
    },
]


# ===================== 报告模板相关 =====================


async def list_templates(db: AsyncSession, report_type: str | None = None) -> list[dict]:
    """列出报告模板"""
    query = select(ReportTemplate)
    if report_type:
        query = query.where(ReportTemplate.report_type == report_type)
    templates = (await db.execute(query)).scalars().all()
    return [
        {
            "id": str(t.id),
            "name": t.name,
            "report_type": t.report_type,
            "sections": t.sections or [],
            "is_default": t.is_default,
            "desc": t.desc,
            "created_at": t.created_at.isoformat() if t.created_at else None,
            "updated_at": t.updated_at.isoformat() if t.updated_at else None,
        }
        for t in templates
    ]


async def get_template_by_id(db: AsyncSession, template_id: str) -> ReportTemplate | None:
    from uuid import UUID
    try:
        uid = UUID(template_id)
    except (ValueError, AttributeError):
        return None
    return (await db.execute(select(ReportTemplate).where(ReportTemplate.id == uid))).scalar_one_or_none()


async def get_default_template(db: AsyncSession, report_type: str) -> ReportTemplate | None:
    """获取指定类型的默认模板（优先 is_default，否则取首个）"""
    t = (await db.execute(
        select(ReportTemplate)
        .where(ReportTemplate.report_type == report_type, ReportTemplate.is_default.is_(True))
        .limit(1)
    )).scalar_one_or_none()
    if t:
        return t
    return (await db.execute(
        select(ReportTemplate).where(ReportTemplate.report_type == report_type).limit(1)
    )).scalar_one_or_none()


async def seed_default_templates(db: AsyncSession) -> int:
    """按 report_type 补齐缺失的默认模板：某类型无任何模板时才写入。
    兼容已有库（仅新增类型会插入默认模板）。返回写入数量。"""
    count = 0
    for data in DEFAULT_TEMPLATES:
        existing = (await db.execute(
            select(ReportTemplate).where(ReportTemplate.report_type == data["report_type"]).limit(1)
        )).scalar_one_or_none()
        if not existing:
            db.add(ReportTemplate(**data))
            count += 1
    if count:
        await db.commit()
    return count


async def create_template(
    db: AsyncSession,
    name: str,
    report_type: str,
    sections: list,
    is_default: bool = False,
    desc: str | None = None,
) -> dict:
    if report_type not in ("antibody_analysis", "vaccination_strategy", "immune_barrier_assessment"):
        raise ValueError("report_type 必须是 antibody_analysis、vaccination_strategy 或 immune_barrier_assessment")
    # 确保 sections 有序且含必要字段
    for i, s in enumerate(sections):
        if "title" not in s or "type" not in s:
            raise ValueError("每个章节必须包含 title 和 type")
        s.setdefault("order", i)
        s.setdefault("content_template", "")
    if is_default:
        # 同类型取消原默认标记
        await db.execute(
            _update_defaults_stmt(report_type)
        )
    t = ReportTemplate(name=name, report_type=report_type, sections=sections, is_default=is_default, desc=desc)
    db.add(t)
    await db.commit()
    await db.refresh(t)
    return _template_to_dict(t)


async def update_template(
    db: AsyncSession,
    template_id: str,
    name: str | None = None,
    sections: list | None = None,
    is_default: bool | None = None,
    desc: str | None = None,
) -> dict | None:
    t = await get_template_by_id(db, template_id)
    if not t:
        return None
    if name is not None:
        t.name = name
    if sections is not None:
        for i, s in enumerate(sections):
            s.setdefault("order", i)
            s.setdefault("content_template", "")
        t.sections = sections
    if desc is not None:
        t.desc = desc
    if is_default is not None:
        if is_default:
            await db.execute(_update_defaults_stmt(t.report_type))
        t.is_default = is_default
    t.updated_at = datetime.now(timezone.utc)
    await db.commit()
    await db.refresh(t)
    return _template_to_dict(t)


async def delete_template(db: AsyncSession, template_id: str) -> bool:
    from sqlalchemy import delete
    result = await db.execute(delete(ReportTemplate).where(ReportTemplate.id == UUID(template_id)))
    await db.commit()
    return result.rowcount > 0


def _update_defaults_stmt(report_type: str):
    from sqlalchemy import update
    return update(ReportTemplate).where(
        ReportTemplate.report_type == report_type, ReportTemplate.is_default.is_(True)
    ).values(is_default=False)


def _template_to_dict(t: ReportTemplate) -> dict:
    return {
        "id": str(t.id),
        "name": t.name,
        "report_type": t.report_type,
        "sections": t.sections or [],
        "is_default": t.is_default,
        "desc": t.desc,
        "created_at": t.created_at.isoformat() if t.created_at else None,
        "updated_at": t.updated_at.isoformat() if t.updated_at else None,
    }


# ===================== 模板章节渲染 =====================


def _build_context(rows: list, disease: str | None = None, language: str = "zh") -> dict:
    """从审核通过数据点构建统一上下文，供各类型章节渲染使用。"""

    lit_ids = {str(r.literature_id) for r in rows if r.literature_id}
    provinces_set = set()
    for r in rows:
        for p in (r.province or "").split(";"):
            p = p.strip()
            if p:
                provinces_set.add(p)
    total_sample = sum(r.sample_size or 0 for r in rows)
    overall_rate = _calc_weighted_rate(rows)[0]

    # 分省
    province_map: dict = {}
    for r in rows:
        for p in (r.province or "").split(";"):
            p = p.strip() or "未知"
            province_map.setdefault(p, []).append(r)
    province_rows = []
    for p, group in sorted(province_map.items()):
        wpr, ps = _calc_weighted_rate(group)
        province_rows.append((p, wpr, ps, len(group)))
    province_rows.sort(key=lambda x: -x[1])

    # 分年
    year_map: dict = {}
    for r in rows:
        if r.collection_year is not None:
            year_map.setdefault(r.collection_year, []).append(r)
    year_rows = []
    for y in sorted(year_map.keys()):
        group = year_map[y]
        wpr, ps = _calc_weighted_rate(group)
        year_rows.append((y, wpr, ps, len(group)))

    # 分年龄
    age_labels = AGE_GROUPS if language == "zh" else AGE_GROUPS_EN
    age_map: dict = {}
    for r in rows:
        matched = False
        for label, lo, hi in age_labels:
            if r.age_min is not None and r.age_max is not None and r.age_min >= lo and r.age_max <= hi:
                age_map.setdefault(label, []).append(r)
                matched = True
                break
        if not matched:
            other_label = "其他" if language == "zh" else "Other"
            age_map.setdefault(other_label, []).append(r)
    age_rows = []
    for label, _lo, _hi in age_labels:
        group = age_map.get(label, [])
        if group:
            wpr, ps = _calc_weighted_rate(group)
            age_rows.append((label, wpr, ps))
    other_label = "其他" if language == "zh" else "Other"
    if age_map.get(other_label):
        group = age_map[other_label]
        wpr, ps = _calc_weighted_rate(group)
        age_rows.append((other_label, wpr, ps))

    # 分疾病（疫苗接种策略场景）
    disease_map: dict = {}
    for r in rows:
        key = r.disease or "未知"
        disease_map.setdefault(key, []).append(r)
    disease_rows = []
    for key, group in disease_map.items():
        wpr, ps = _calc_weighted_rate(group)
        gmc_rows = [r for r in group if r.data_type == "gmc" and r.value is not None]
        avg_gmc = round(sum(r.value for r in gmc_rows) / len(gmc_rows), 2) if gmc_rows else None
        disease_rows.append((DISEASE_NAMES.get(key, key), wpr, ps, len(group), avg_gmc))
    disease_rows.sort(key=lambda x: -x[1])

    return {
        "disease_name": disease,
        "language": language,
        "literature_count": len(lit_ids),
        "point_count": len(rows),
        "province_count": len(provinces_set),
        "total_samples": total_sample,
        "weighted_rate": overall_rate,
        "province_rows": province_rows,
        "year_rows": year_rows,
        "age_rows": age_rows,
        "disease_rows": disease_rows,
    }


def _fmt_rate(v) -> str:
    return f"{v}%" if v is not None else "-"


def _render_kpi(section: dict, ctx: dict) -> str:
    keys = section.get("kpi") or []
    labels = {
        "literature_count": ("文献数", lambda c: str(c)),
        "point_count": ("数据点", lambda c: str(c)),
        "province_count": ("覆盖省份", lambda c: str(c)),
        "total_samples": ("样本量", lambda c: str(c)),
        "weighted_rate": ("加权阳性率", lambda c: _fmt_rate(c)),
    }
    lines = []
    for k in keys:
        if k not in labels or k not in ctx:
            continue
        label, fmt = labels[k]
        lines.append(f"  - **{label}**：{fmt(ctx[k])}")
    if not lines:
        return "暂无关键指标数据。"
    return "\n".join(lines)


def _render_table(section: dict, ctx: dict) -> str:
    data = section.get("data", "province")
    if data == "year":
        rows, headers = ctx["year_rows"], ("年份", "加权阳性率", "样本量", "数据点")
    elif data == "age":
        rows, headers = ctx["age_rows"], ("年龄段", "加权阳性率", "样本量")
    elif data == "disease":
        rows, headers = ctx["disease_rows"], ("疾病", "加权阳性率", "样本量", "数据点", "GMC")
    else:
        rows, headers = ctx["province_rows"], ("省份", "加权阳性率", "样本量", "数据点")

    if not rows:
        return "暂无该维度数据。"

    def _cell(v):
        return "-" if v is None else str(v)

    body = "\n".join(
        "| " + " | ".join(_cell(c if i == 0 else (round(c, 2) if isinstance(c, float) else c)) if i == 0 else _cell(c) for i, c in enumerate(r)) + " |"
        for r in rows
    )
    header = "| " + " | ".join(headers) + " |"
    sep = "| " + " | ".join(["---"] * len(headers)) + " |"
    return f"{header}\n{sep}\n{body}"


def _render_chart(section: dict, ctx: dict) -> str:
    """将 chart 章节渲染为「真实图表（内联 SVG）+ 摘要注记 + 明细表」。

    图表以 data URI 形式内联进 Markdown，无需外部图片资源、可随报告内容存库，
    在支持 Markdown 图片的前端（react-markdown）中直接渲染。
    """
    analysis = section.get("analysis", "trend")
    lang = str(ctx.get("language", "zh"))
    no_data = "暂无该维度数据。" if lang.startswith("zh") else "No data available for this dimension."

    if analysis == "region":
        items = [(r[0], r[1]) for r in ctx["province_rows"]]
        title = "地区分布（加权阳性率）" if lang.startswith("zh") else "Provincial distribution (weighted positivity)"
        table = _render_table({**section, "data": "province"}, ctx)
    elif analysis == "age_curve":
        items = [(r[0], r[1]) for r in ctx["age_rows"]]
        title = "年龄分布（加权阳性率）" if lang.startswith("zh") else "Age-stratified weighted positivity"
        table = _render_table({**section, "data": "age"}, ctx)
    elif analysis == "disease":
        items = [(r[0], r[1]) for r in ctx["disease_rows"]]
        title = "疾病流行（加权阳性率）" if lang.startswith("zh") else "Disease-wise weighted positivity"
        table = _render_table({**section, "data": "disease"}, ctx)
    else:  # trend
        items = [(r[0], r[1]) for r in ctx["year_rows"]]
        title = "时间趋势（加权阳性率）" if lang.startswith("zh") else "Time trend (weighted positivity)"
        table = _render_table({**section, "data": "year"}, ctx)

    if not items:
        return no_data

    # 图片仅展示 Top k（防止地区/疾病过多导致 X 轴拥挤），明细表仍保留全部
    top_items = _chart_top_items(items, 16)
    note = _chart_summary_note(top_items, analysis, lang)
    img_md = _chart_data_uri("bar" if analysis != "trend" else "line",
                             title, top_items, lang)
    return f"{img_md}\n\n{note}\n\n{table}"


# ---------------------------------------------------------------------------
# 报告内联图表（依赖无关的纯 SVG 生成，避免容器内缺少中文字库导致栅格化乱码）
# ---------------------------------------------------------------------------

def _xml_escape(s) -> str:
    return (str(s).replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


def _chart_top_items(items, k: int):
    clean = [it for it in items if it[1] is not None]
    clean.sort(key=lambda it: it[1], reverse=True)
    return clean[:k]


def _chart_summary_note(items, analysis: str, lang: str) -> str:
    values = [v for _, v in items if v is not None]
    if lang.startswith("zh"):
        if analysis == "trend":
            head = f"共 {len(values)} 个年份"
        elif analysis == "age_curve":
            head = f"共 {len(values)} 个年龄段"
        elif analysis == "region":
            head = f"共覆盖 {len(values)} 个省份/地区"
        else:
            head = f"共 {len(values)} 类疾病"
        note = head + (f"，加权阳性率区间 {min(values):g}% ~ {max(values):g}%" if values else "")
        if len(values) < len(items):
            note += f"（前 {len(values)} 名）"
        return f"根据图表数据：{note}。"
    tail = f"，区间 {min(values):g}% ~ {max(values):g}%" if values else ""
    return f"Relevant data: {len(values)} categories{tail}."


def _chart_sort_key(label) -> str:
    """时间趋势按数值年份排序（兼容 '2020' 与 2020 两种形式）。"""
    try:
        return f"{int(label):08d}"
    except (TypeError, ValueError):
        return str(label)


def _chart_svg(kind: str, title: str, items, lang: str = "zh") -> str:
    """生成纯 SVG 图表（bar/line），坐标标签已做 XML 转义。

    items: list[(label, value)]；value 为百分数（0-100），None 已在上游剔除。
    """
    data = [(_xml_escape(str(label)), float(v)) for label, v in items if v is not None]
    W, H = 720, 380
    ml, mr, mt, mb = 78, 26, 52, 78
    pw, ph = W - ml - mr, H - mt - mb
    empty_txt = "无数据" if lang.startswith("zh") else "No data"

    if not data:
        return (f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" '
                f'viewBox="0 0 {W} {H}"><text x="{W / 2}" y="{H / 2}" text-anchor="middle" '
                f'font-size="16" fill="#9aa5b1">{empty_txt}</text></svg>')

    if kind == "line":
        data.sort(key=lambda t: _chart_sort_key(t[0]))
    labels = [d[0] for d in data]
    vals = [d[1] for d in data]
    n = len(data)

    ymax = max(0.0, max(vals))
    ytop = (ymax * 1.15) if ymax > 0 else 1.0

    def sx(i):
        if n == 1:
            return ml + pw / 2
        return ml + pw * i / (n - 1)

    def sy(v):
        return mt + ph * (1 - v / ytop)

    parts = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" viewBox="0 0 {W} {H}">']
    parts.append(f'<text x="{ml}" y="{mt - 20}" font-size="16" font-weight="bold" fill="#333333">{_xml_escape(title)}</text>')

    for k in range(5):
        vv = ytop * k / 4
        yy = sy(vv)
        parts.append(f'<line x1="{ml}" y1="{yy:.1f}" x2="{W - mr}" y2="{yy:.1f}" stroke="#eaeef2" stroke-width="1"/>')
        parts.append(f'<text x="{ml - 8}" y="{yy + 4:.1f}" text-anchor="end" font-size="11" fill="#6b7280">{vv:g}</text>')
    parts.append(f'<line x1="{ml}" y1="{mt}" x2="{ml}" y2="{mt + ph}" stroke="#94a3b8" stroke-width="1.5"/>')
    parts.append(f'<line x1="{ml}" y1="{mt + ph}" x2="{W - mr}" y2="{mt + ph}" stroke="#94a3b8" stroke-width="1.5"/>')

    if kind == "line":
        pts = " ".join(f"{sx(i):.1f},{sy(v):.1f}" for i, v in enumerate(vals))
        area = f"{sx(0):.1f},{mt + ph:.1f} {pts} {sx(n - 1):.1f},{mt + ph:.1f}"
        parts.append(f'<polygon points="{area}" fill="#4e79a7" fill-opacity="0.10"/>')
        parts.append(f'<polyline points="{pts}" fill="none" stroke="#4e79a7" stroke-width="2.5" stroke-linejoin="round" stroke-linecap="round"/>')
        for i, v in enumerate(vals):
            parts.append(f'<circle cx="{sx(i):.1f}" cy="{sy(v):.1f}" r="3.5" fill="#4e79a7"/>')
        for i, lab in enumerate(labels):
            parts.append(f'<text x="{sx(i):.1f}" y="{mt + ph + 18}" text-anchor="middle" font-size="12" fill="#374151">{lab}</text>')
    else:
        bar_w = min(46.0, pw / max(n, 1) * 0.62)
        for i, v in enumerate(vals):
            x0 = sx(i) - bar_w / 2
            y0 = sy(v)
            parts.append(f'<rect x="{x0:.1f}" y="{y0:.1f}" width="{bar_w:.1f}" height="{max(mt + ph - y0, 0):.1f}" rx="2" fill="#4e79a7"/>')
            parts.append(f'<text x="{sx(i):.1f}" y="{y0 - 6:.1f}" text-anchor="middle" font-size="11" fill="#374151">{v:g}%</text>')
        for i, lab in enumerate(labels):
            parts.append(f'<text x="{sx(i):.1f}" y="{mt + ph + 34}" text-anchor="end" '
                         f'transform="rotate(-45 {sx(i):.1f} {mt + ph + 34})" font-size="11" fill="#374151">{lab}</text>')
    parts.append("</svg>")
    return "\n".join(parts)


def _chart_data_uri(kind: str, title: str, items, lang: str = "zh") -> str:
    """生成可直接内联进 Markdown 的图表图片行：![…](data:image/svg+xml;base64,…)。"""
    svg = _chart_svg(kind, title, items, lang).encode("utf-8")
    return f"![{_xml_escape(title)}](data:image/svg+xml;base64,{base64.b64encode(svg).decode('ascii')})"


def _context_text(ctx: dict) -> str:
    """把上下文汇总为纯文本，供 LLM 章节提示使用。"""
    lines = [
        f"- 数据来源：{ctx['literature_count']} 篇文献，{ctx['point_count']} 个数据点",
        f"- 覆盖省份：{ctx['province_count']} 个；总样本量：{ctx['total_samples']} 人；加权阳性率：{_fmt_rate(ctx['weighted_rate'])}",
    ]
    lines.append("\n分省：")
    lines += [f"  - {p}：阳性率 {wpr}%，样本量 {ps}，数据点 {n} 个" for p, wpr, ps, n in ctx["province_rows"]]
    if ctx["year_rows"]:
        lines.append("\n分年：")
        lines += [f"  - {y}年：阳性率 {wpr}%，样本量 {ps}" for y, wpr, ps, _ in ctx["year_rows"]]
    if ctx["age_rows"]:
        lines.append("\n分年龄：")
        lines += [f"  - {label}：阳性率 {wpr}%，样本量 {ps}" for label, wpr, ps in ctx["age_rows"]]
    if ctx["disease_rows"] and len(ctx["disease_rows"]) > 1:
        lines.append("\n分疾病：")
        lines += [f"  - {name}：阳性率 {wpr}%，样本量 {ps}，数据点 {n} 个" for name, wpr, ps, n, _ in ctx["disease_rows"]]
    return "\n".join(lines)


async def _render_template_report(
    db: AsyncSession,
    ctx: dict,
    template: ReportTemplate,
    model: str | None,
    language: str,
) -> str:
    """按模板 sections 顺序渲染报告 Markdown。"""
    blocks = []
    context_text = _context_text(ctx)
    sections = sorted(template.sections or [], key=lambda s: s.get("order", 0))
    for s in sections:
        stype = s.get("type", "text")
        title = s.get("title", "章节")
        if stype == "kpi":
            body = _render_kpi(s, ctx)
        elif stype == "table":
            body = _render_table(s, ctx)
        elif stype == "chart":
            body = _render_chart(s, ctx)
        else:  # text
            instruction = s.get("content_template") or "请基于以下数据撰写本章节的分析内容。"
            prompt = (
                f"你是一位流行病学专家。请撰写报告章节《{title}》。\n"
                f"章节要求：{instruction}\n\n"
                f"可用数据：\n{context_text}\n\n"
                f"请直接输出该章节的正文内容（Markdown），不要输出章节标题，不要编造不存在的数据。"
            )
            body = await _call_llm(db, prompt, model=model)
        blocks.append(f"## {title}\n\n{body.strip()}")
    return "\n\n".join(blocks)


def _build_legacy_inline_text(rows: list, language: str = "zh") -> tuple[str, str, str]:
    """按内置 Prompt 所需的文本格式，构建分省/分年/分年龄摘要（兜底路径）。"""
    province_lines = []
    province_map: dict[str, list[DataPoint]] = {}
    for r in rows:
        for p in (r.province or "").split(";"):
            p = p.strip() or "未知"
            province_map.setdefault(p, []).append(r)
    for prov, group in sorted(province_map.items()):
        wpr, ps = _calc_weighted_rate(group)
        province_lines.append(f"- {prov}：阳性率 {wpr}%，样本量 {ps}，数据点 {len(group)} 个")
    province_table = "\n".join(province_lines) if province_lines else "暂无数据"

    year_map: dict[int, list[DataPoint]] = {}
    for r in rows:
        y = r.collection_year
        if y is None:
            continue
        year_map.setdefault(y, []).append(r)
    year_lines = []
    for year in sorted(year_map.keys()):
        group = year_map[year]
        wpr, ps = _calc_weighted_rate(group)
        year_lines.append(f"- {year}年：阳性率 {wpr}%，样本量 {ps}，数据点 {len(group)} 个")
    year_trend = "\n".join(year_lines) if year_lines else "暂无数据"

    age_labels = AGE_GROUPS if language == "zh" else AGE_GROUPS_EN
    age_map: dict[str, list[DataPoint]] = {label: [] for label, _, _ in age_labels}
    other_label = "其他" if language == "zh" else "Other"
    age_map.setdefault(other_label, [])
    for r in rows:
        matched = False
        for label, lo, hi in age_labels:
            if r.age_min is not None and r.age_max is not None and r.age_min >= lo and r.age_max <= hi:
                age_map[label].append(r)
                matched = True
                break
        if not matched:
            age_map[other_label].append(r)
    age_lines = []
    for label, _lo, _hi in age_labels:
        group = age_map.get(label, [])
        if group:
            wpr, ps = _calc_weighted_rate(group)
            age_lines.append(f"- {label}：阳性率 {wpr}%，样本量 {ps}")
    if age_map.get(other_label):
        group = age_map[other_label]
        wpr, ps = _calc_weighted_rate(group)
        age_lines.append(f"- {other_label}：阳性率 {wpr}%，样本量 {ps}")
    age_distribution = "\n".join(age_lines) if age_lines else "暂无数据"

    return province_table, year_trend, age_distribution


async def generate_report(
    db: AsyncSession,
    disease: str | None = None,
    province: str | None = None,
    data_type: str | None = None,
    language: str = "zh",
    title: str | None = None,
    model: str | None = None,
    template_id: str | None = None,
) -> dict:
    """生成报告。

    若提供 template_id，则按模板的 sections 顺序渲染（支持 text/chart/table/kpi 章节）；
    否则使用内置默认 Prompt 结构生成。
    """
    # 0. 解析模板（template_id 缺省时使用抗体分析默认模板）
    template = None
    if template_id:
        template = await get_template_by_id(db, template_id)
        if not template:
            raise ValueError("指定的报告模板不存在")
    else:
        template = await get_default_template(db, "antibody_analysis")

    # 1. 查询审核通过的数据点
    query = select(DataPoint).where(DataPoint.review_status == "approved")
    if disease:
        query = query.where(DataPoint.disease == disease)
    if province:
        query = query.where(DataPoint.province.ilike(f"%{province}%"))
    if data_type:
        query = query.where(DataPoint.data_type == data_type)

    result = await db.execute(query)
    rows = list(result.scalars().all())

    if not rows:
        raise ValueError("没有找到审核通过的数据，无法生成报告")

    # 释放读事务占用的数据库连接：本地 LLM 推理可能耗时数分钟，若连接一直处于
    # idle-in-transaction，会被 PostgreSQL 的 idle_in_transaction_session_timeout
    # （120s）强制断开，导致后续保存报告时出现
    # "cannot call Transaction.rollback(): the underlying connection is closed"。
    # rows 已整体加载进内存且 expire_on_commit=False，commit 后数据仍可用。
    await db.commit()

    # 2. 生成疾病名称与报告标题
    disease_name = DISEASE_NAMES.get(disease or "", disease or "未知疾病")

    # 引用/编号统一固定为生成当日时间（一次性计算），避免多处 date.today() 不一致
    gen_date = datetime.now(timezone.utc)

    if title:
        report_title = title
        report_title_en = title
    else:
        report_title = f"{disease_name}抗体水平分析报告"
        report_title_en = f"{disease_name} Antibody Level Analysis Report"

    # 8. 生成正文：template_id 走模板渲染，否则走内置 Prompt
    if template:
        ctx = _build_context(rows, disease=disease, language=language)
        content = await _render_template_report(db, ctx, template, model=model, language=language)
    else:
        content = None

    if content is None:
        # 使用内置 Prompt（无可用模板时的兜底路径）
        lit_ids = {str(r.literature_id) for r in rows if r.literature_id}
        provinces_set = set()
        for r in rows:
            for p in (r.province or "").split(";"):
                p = p.strip()
                if p:
                    provinces_set.add(p)
        total_sample = sum(r.sample_size or 0 for r in rows)
        province_table, year_trend, age_distribution = _build_legacy_inline_text(rows, language)
        if language == "zh":
            prompt = REPORT_PROMPT_ZH.format(
                title=report_title,
                literature_count=len(lit_ids),
                point_count=len(rows),
                province_count=len(provinces_set),
                total_samples=total_sample,
                province_table=province_table,
                year_trend=year_trend,
                age_distribution=age_distribution,
            )
        else:
            prompt = REPORT_PROMPT_EN.format(
                title_en=report_title_en,
                literature_count=len(lit_ids),
                point_count=len(rows),
                province_count=len(provinces_set),
                total_samples=total_sample,
                province_table=province_table,
                year_trend=year_trend,
                age_distribution=age_distribution,
            )
        content = await _call_llm(db, prompt, model=model)

    # 8.5 方法学小节：统一脚注（复用 build_methodology_note），拼接进报告正文
    lit_ids = {str(r.literature_id) for r in rows if r.literature_id}
    methodology_note = build_methodology_note(
        "report",
        {"disease": disease, "province": province, "data_type": data_type},
        {"n_estimates": len(rows), "n_literatures": len(lit_ids), "quality_grades": True},
    )
    content = (content or "").rstrip()
    if language == "zh":
        content += f"\n\n## 方法学\n\n{methodology_note}"
        content += (
            f"\n\n## 引用\n\n"
            f"抗体地图数据库分析报告[EB/OL]. 抗体地图数据库（版本 v1.0）. "
            f"数据截至：{gen_date.date().isoformat()}；[引用日期 {gen_date.date().isoformat()}]. "
            f"报告编号：{disease_name}_{province or '全国'}_{gen_date.date().isoformat()}。"
        )
    else:
        content += f"\n\n## Methodology\n\n{methodology_note}"
        content += (
            f"\n\n## Citation\n\n"
            f"Antibody Map Database Analysis Report[EB/OL]. Antibody Map Database (Version v1.0). "
            f"Data as of: {gen_date.date().isoformat()}; [Accessed {gen_date.date().isoformat()}]. "
            f"Report ID: {disease_name}_{province or 'National'}_{gen_date.date().isoformat()}."
        )

    # 解析模型显示名称
    llm_model_name = await _resolve_model_name(db, model)

    # 9. Save to database
    try:
        report = Report(
            title=report_title,
            content=content,
            report_type="antibody_analysis",
            disease=disease,
            province=province,
            data_type=data_type,
            language=language,
            literature_count=len(lit_ids),
            data_point_count=len(rows),
            llm_model=llm_model_name,
            data_snapshot_hash=_data_snapshot_hash(rows),
            generated_at=gen_date,
        )
        db.add(report)
        await db.commit()
    except Exception as e:
        logger.error(f"保存抗体分析报告失败: {e}；内容摘要: {content[:300]}")
        raise RuntimeError(f"报告生成成功但保存失败: {e}") from e

    return {
        "id": str(report.id),
        "title": report_title,
        "content": content,
        "report_type": "antibody_analysis",
        "literature_count": len(lit_ids),
        "data_point_count": len(rows),
        "language": language,
        "llm_model": report.llm_model,
        "generated_at": report.generated_at.isoformat(),
    }


async def generate_immune_barrier_report(
    db: AsyncSession,
    disease: str | None = None,
    province: str | None = None,
    data_type: str | None = None,
    language: str = "zh",
    title: str | None = None,
    model: str | None = None,
    template_id: str | None = None,
) -> dict:
    """生成免疫屏障评估报告。

    参数结构与抗体分析报告一致，复用同一批审核通过数据点，
    但输出侧重人群免疫屏障总体/地区/时间/年龄维度评估与缺口干预建议。
    """
    # 0. 解析模板（缺省时使用免疫屏障评估默认模板）
    template = None
    if template_id:
        template = await get_template_by_id(db, template_id)
        if not template:
            raise ValueError("指定的报告模板不存在")
    else:
        template = await get_default_template(db, "immune_barrier_assessment")

    # 1. 查询审核通过的数据点
    query = select(DataPoint).where(DataPoint.review_status == "approved")
    if disease:
        query = query.where(DataPoint.disease == disease)
    if province:
        query = query.where(DataPoint.province.ilike(f"%{province}%"))
    if data_type:
        query = query.where(DataPoint.data_type == data_type)

    result = await db.execute(query)
    rows = list(result.scalars().all())

    if not rows:
        raise ValueError("没有找到审核通过的数据，无法生成报告")

    # 释放读事务占用的数据库连接（同 generate_report）：本地 LLM 推理耗时数分钟，
    # 避免连接因 idle_in_transaction_session_timeout 被断开导致保存失败。
    await db.commit()

    disease_name = DISEASE_NAMES.get(disease or "", disease or "未知疾病")

    # 引用/编号统一固定为生成当日时间（一次性计算），避免多处 date.today() 不一致
    gen_date = datetime.now(timezone.utc)

    if title:
        report_title = title
        report_title_en = title
    else:
        report_title = f"{disease_name}免疫屏障评估报告"
        report_title_en = f"{disease_name} Immune Barrier Assessment Report"

    # 2. 模板渲染或内置 Prompt
    if template:
        ctx = _build_context(rows, disease=disease, language=language)
        content = await _render_template_report(db, ctx, template, model=model, language=language)
    else:
        lit_ids = {str(r.literature_id) for r in rows if r.literature_id}
        provinces_set = set()
        for r in rows:
            for p in (r.province or "").split(";"):
                p = p.strip()
                if p:
                    provinces_set.add(p)
        total_sample = sum(r.sample_size or 0 for r in rows)
        province_table, year_trend, age_distribution = _build_legacy_inline_text(rows, language)
        if language == "zh":
            prompt = IMMUNE_BARRIER_PROMPT_ZH.format(
                title=report_title,
                literature_count=len(lit_ids),
                point_count=len(rows),
                province_count=len(provinces_set),
                total_samples=total_sample,
                province_table=province_table,
                year_trend=year_trend,
                age_distribution=age_distribution,
            )
        else:
            prompt = IMMUNE_BARRIER_PROMPT_EN.format(
                title_en=report_title_en,
                literature_count=len(lit_ids),
                point_count=len(rows),
                province_count=len(provinces_set),
                total_samples=total_sample,
                province_table=province_table,
                year_trend=year_trend,
                age_distribution=age_distribution,
            )
        content = await _call_llm(db, prompt, model=model)

    lit_ids = {str(r.literature_id) for r in rows if r.literature_id}
    methodology_note = build_methodology_note(
        "report",
        {"disease": disease, "province": province, "data_type": data_type},
        {"n_estimates": len(rows), "n_literatures": len(lit_ids), "quality_grades": True},
    )
    content = (content or "").rstrip()
    if language == "zh":
        content += f"\n\n## 方法学\n\n{methodology_note}"
        content += (
            f"\n\n## 引用\n\n"
            f"抗体地图数据库分析报告[EB/OL]. 抗体地图数据库（版本 v1.0）. "
            f"数据截至：{gen_date.date().isoformat()}；[引用日期 {gen_date.date().isoformat()}]. "
            f"报告编号：{disease_name}_{province or '全国'}_{gen_date.date().isoformat()}。"
        )
    else:
        content += f"\n\n## Methodology\n\n{methodology_note}"
        content += (
            f"\n\n## Citation\n\n"
            f"Antibody Map Database Analysis Report[EB/OL]. Antibody Map Database (Version v1.0). "
            f"Data as of: {gen_date.date().isoformat()}; [Accessed {gen_date.date().isoformat()}]. "
            f"Report ID: {disease_name}_{province or 'National'}_{gen_date.date().isoformat()}."
        )

    llm_model_name = await _resolve_model_name(db, model)

    try:
        report = Report(
            title=report_title,
            content=content,
            report_type="immune_barrier_assessment",
            disease=disease,
            province=province,
            data_type=data_type,
            language=language,
            literature_count=len(lit_ids),
            data_point_count=len(rows),
            llm_model=llm_model_name,
            data_snapshot_hash=_data_snapshot_hash(rows),
            generated_at=gen_date,
        )
        db.add(report)
        await db.commit()
    except Exception as e:
        logger.error(f"保存免疫屏障评估报告失败: {e}；内容摘要: {content[:300]}")
        raise RuntimeError(f"报告生成成功但保存失败: {e}") from e

    return {
        "id": str(report.id),
        "title": report_title,
        "content": content,
        "report_type": "immune_barrier_assessment",
        "literature_count": len(lit_ids),
        "data_point_count": len(rows),
        "language": language,
        "llm_model": report.llm_model,
        "generated_at": report.generated_at.isoformat(),
    }


async def generate_vaccination_strategy_report(
    db: AsyncSession,
    task_type: str,
    task_time: str,
    task_location: str,
    personnel_count: int,
    personnel_gender: str = "",
    personnel_age: str = "",
    personnel_vaccination_history: str = "",
    title: str | None = None,
    template_id: str | None = None,
    model: str | None = None,
) -> dict:
    """生成疫苗接种策略研判报告。template_id 缺省时使用内置 Prompt；指定时按模板 sections 渲染。"""
    # 1. 查询任务地点的传染病流行数据
    query = select(DataPoint).where(DataPoint.review_status == "approved")
    if task_location:
        query = query.where(DataPoint.province.ilike(f"%{task_location}%"))

    result = await db.execute(query)
    rows = list(result.scalars().all())

    # 释放读事务占用的数据库连接（同 generate_report）：本地 LLM 推理耗时数分钟，
    # 避免连接因 idle_in_transaction_session_timeout 被断开导致保存失败。
    await db.commit()

    # 2. 构建疫情数据汇总
    epidemic_lines = []
    if rows:
        lit_ids = {str(r.literature_id) for r in rows if r.literature_id}
        epidemic_lines.append(f"- 数据来源：{len(lit_ids)} 篇文献，{len(rows)} 个数据点")

        # 按疾病汇总
        disease_map: dict[str, list[DataPoint]] = {}
        for r in rows:
            d = r.disease or "未知"
            if d not in disease_map:
                disease_map[d] = []
            disease_map[d].append(r)

        epidemic_lines.append("\n各疾病流行情况：")
        for disease_key, group in disease_map.items():
            disease_name = DISEASE_NAMES.get(disease_key, disease_key)
            wpr, ps = _calc_weighted_rate(group)
            gmc_rows = [r for r in group if r.data_type == "gmc" and r.value is not None]
            avg_gmc = round(sum(r.value for r in gmc_rows) / len(gmc_rows), 2) if gmc_rows else None
            epidemic_lines.append(
                f"  - {disease_name}：加权阳性率 {wpr}%（样本量 {ps}）"
                + (f"，GMC {avg_gmc}" if avg_gmc else "")
                + f"，数据点 {len(group)} 个"
            )

        # 按年份汇总
        year_map: dict[int, list[DataPoint]] = {}
        for r in rows:
            y = r.collection_year
            if y is None:
                continue
            if y not in year_map:
                year_map[y] = []
            year_map[y].append(r)

        if year_map:
            epidemic_lines.append("\n年份趋势：")
            for year in sorted(year_map.keys(), reverse=True)[:5]:
                group = year_map[year]
                wpr, ps = _calc_weighted_rate(group)
                epidemic_lines.append(f"  - {year}年：阳性率 {wpr}%，样本量 {ps}，数据点 {len(group)} 个")

        # 按年龄汇总
        age_map: dict[str, list[DataPoint]] = {}
        for r in rows:
            matched = False
            for label, lo, hi in AGE_GROUPS:
                if (
                    r.age_min is not None
                    and r.age_max is not None
                    and r.age_min >= lo
                    and r.age_max <= hi
                ):
                    if label not in age_map:
                        age_map[label] = []
                    age_map[label].append(r)
                    matched = True
                    break
            if not matched:
                if "其他" not in age_map:
                    age_map["其他"] = []
                age_map["其他"].append(r)

        if age_map:
            epidemic_lines.append("\n年龄分布：")
            for label in age_map:
                group = age_map[label]
                wpr, ps = _calc_weighted_rate(group)
                epidemic_lines.append(f"  - {label}：阳性率 {wpr}%，样本量 {ps}")
    else:
        epidemic_lines.append("暂无该地区的传染病流行病学数据。")

    epidemic_data = "\n".join(epidemic_lines)

    # 3. 报告标题
    report_title = title or f"{task_location}任务人员疫苗接种策略研判报告"

    # 3.5 若指定模板，按模板 sections 渲染
    template = None
    if template_id:
        template = await get_template_by_id(db, template_id)
        if not template:
            raise ValueError("指定的报告模板不存在")

    # 4. 构建 Prompt 或渲染模板
    if template:
        ctx = _build_context(rows, language="zh")
        ctx["task_info"] = (
            f"任务类型：{task_type}；任务时间：{task_time}；任务地点：{task_location}；"
            f"人员人数：{personnel_count}人；年龄范围：{personnel_age or '未提供'}；"
            f"疫苗接种史：{personnel_vaccination_history or '未提供'}"
        )
        content = await _render_template_report(db, ctx, template, model=model, language="zh")
    else:
        prompt = VACCINATION_STRATEGY_PROMPT_ZH.format(
            task_type=task_type,
            task_time=task_time,
            task_location=task_location,
            personnel_count=personnel_count,
            personnel_gender=personnel_gender or "未提供",
            personnel_age=personnel_age or "未提供",
            personnel_vaccination_history=personnel_vaccination_history or "未提供",
            epidemic_data=epidemic_data,
        )
        content = await _call_llm(db, prompt, model=model)

    # 解析模型显示名称
    llm_model_name = await _resolve_model_name(db, model)

    # 5. Save to database
    try:
        report = Report(
            title=report_title,
            content=content,
            report_type="vaccination_strategy",
            task_type=task_type,
            task_time=task_time,
            task_location=task_location,
            personnel_count=personnel_count,
            personnel_gender=personnel_gender or None,
            personnel_age=personnel_age or None,
            personnel_vaccination_history=personnel_vaccination_history or None,
            data_point_count=len(rows),
            literature_count=len({str(r.literature_id) for r in rows if r.literature_id}),
            llm_model=llm_model_name,
        )
        db.add(report)
        await db.commit()
    except Exception as e:
        logger.error(f"保存疫苗接种策略报告失败: {e}；内容摘要: {content[:300]}")
        raise RuntimeError(f"报告生成成功但保存失败: {e}") from e

    return {
        "id": str(report.id),
        "title": report_title,
        "content": content,
        "report_type": "vaccination_strategy",
        "task_type": task_type,
        "task_time": task_time,
        "task_location": task_location,
        "personnel_count": personnel_count,
        "personnel_gender": personnel_gender,
        "personnel_age": personnel_age,
        "personnel_vaccination_history": personnel_vaccination_history,
        "data_point_count": len(rows),
        "literature_count": len({str(r.literature_id) for r in rows if r.literature_id}),
        "llm_model": llm_model_name,
        "language": "zh",
        "generated_at": report.generated_at.isoformat(),
    }


async def _call_llm(db: AsyncSession, prompt: str, model: str | None = None) -> str:
    """调用 LLM 生成报告"""
    llm_model = model or settings.LLM_MODEL
    api_key = settings.LLM_API_KEY
    base_url = settings.LLM_BASE_URL

    # 如果 model 是远程模型配置的 UUID，查找对应的 API 配置
    if model:
        try:
            uid = UUID(model)
            result = await db.execute(select(ApiModelConfig).where(ApiModelConfig.id == uid))
            config = result.scalar_one_or_none()
            if config:
                llm_model = config.model_name
                api_key = config.api_key
                base_url = config.base_url
        except (ValueError, AttributeError):
            # 不是 UUID，当作普通模型名直接使用
            pass

    client = AsyncOpenAI(
        api_key=api_key,
        base_url=base_url,
        timeout=settings.LLM_REQUEST_TIMEOUT,
    )
    try:
        extra_kwargs = {}
        # 本地 Ollama 模型：禁用 thinking 模式
        if ":" in llm_model or llm_model.startswith("qwen"):
            extra_kwargs["extra_body"] = {"think": False}
        response = await client.chat.completions.create(
            model=llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            timeout=settings.LLM_REQUEST_TIMEOUT,
            **extra_kwargs,
        )
        return response.choices[0].message.content or ""
    except Exception as e:
        logger.warning(f"LLM API 调用失败 (model={llm_model}): {e}", exc_info=True)
        raise RuntimeError(f"报告生成失败: {e}") from e


async def get_reports(db: AsyncSession, page: int = 1, page_size: int = 20):
    """??????"""
    from sqlalchemy import desc
    q = select(Report).order_by(desc(Report.generated_at))
    total_q = select(func.count(Report.id))
    total = (await db.execute(total_q)).scalar() or 0
    rows = (await db.execute(q.offset((page - 1) * page_size).limit(page_size))).scalars().all()
    items = []
    for r in rows:
        items.append({
            "id": str(r.id),
            "title": r.title,
            "report_type": r.report_type,
            "disease": r.disease,
            "province": r.province,
            "data_type": r.data_type,
            "language": r.language,
            "llm_model": r.llm_model,
            "literature_count": r.literature_count,
            "data_point_count": r.data_point_count,
            "task_type": r.task_type,
            "task_time": r.task_time,
            "task_location": r.task_location,
            "personnel_count": r.personnel_count,
            "generated_at": r.generated_at.isoformat(),
        })
    return {"items": items, "total": total, "page": page, "page_size": page_size}


async def get_report_by_id(db: AsyncSession, report_id):
    """????????"""
    from uuid import UUID
    uid = UUID(report_id)
    r = (await db.execute(select(Report).where(Report.id == uid))).scalar_one_or_none()
    if not r:
        return None
    return {
        "id": str(r.id),
        "title": r.title,
        "content": r.content,
        "report_type": r.report_type,
        "disease": r.disease,
        "province": r.province,
        "data_type": r.data_type,
        "language": r.language,
        "llm_model": r.llm_model,
        "literature_count": r.literature_count,
        "data_point_count": r.data_point_count,
        "task_type": r.task_type,
        "task_time": r.task_time,
        "task_location": r.task_location,
        "personnel_count": r.personnel_count,
        "personnel_gender": r.personnel_gender,
        "personnel_age": r.personnel_age,
        "personnel_vaccination_history": r.personnel_vaccination_history,
        "data_snapshot_hash": r.data_snapshot_hash,
        "generated_at": r.generated_at.isoformat(),
    }


async def update_report(
    db: AsyncSession,
    report_id: str,
    title: str | None = None,
    content: str | None = None,
) -> dict | None:
    """更新报告标题或内容"""
    from datetime import datetime, timezone
    from uuid import UUID

    from sqlalchemy import update

    uid = UUID(report_id)
    values = {}
    if title is not None:
        values["title"] = title
    if content is not None:
        values["content"] = content
    if not values:
        return await get_report_by_id(db, report_id)

    values["generated_at"] = datetime.now(timezone.utc)
    stmt = update(Report).where(Report.id == uid).values(**values)
    await db.execute(stmt)
    await db.commit()

    return await get_report_by_id(db, report_id)


async def delete_report(db: AsyncSession, report_id: str) -> bool:
    """删除报告"""
    from uuid import UUID

    from sqlalchemy import delete

    uid = UUID(report_id)
    result = await db.execute(delete(Report).where(Report.id == uid))
    await db.commit()
    return result.rowcount > 0


def report_markdown_to_docx(markdown: str) -> bytes:
    """将 Markdown 报告内容转换为 Word（.docx）二进制。

    支持：#/##/### 标题、普通段落、-/* 无序列表、1. 有序列表、
    | 表格、**加粗**、行内代码、``` 代码块、--- 分隔线。
    """
    import re
    from io import BytesIO

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def _add_runs(par, text: str) -> None:
        """解析 **加粗** 与 `行内代码`，逐段添加 run。"""
        for tok in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
                run = par.add_run(tok[2:-2])
                run.bold = True
            elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
                run = par.add_run(tok[1:-1])
                run.font.name = "Consolas"
            else:
                par.add_run(tok)

    def _add_hr(par) -> None:
        """在段落下方添加一条水平线（代替 Markdown 分隔线）。"""
        p_pr = par._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    in_code_block = False
    active_table = None
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        # 跳过内联图表（SVG data URI）：python-docx 无法嵌入 SVG，明细表已在正文保留
        if re.match(r"^!\[.*?\]\(data:image/svg\+xml;base64,", stripped):
            active_table = None
            continue

        # 代码块
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            active_table = None
            continue
        if in_code_block:
            par = doc.add_paragraph()
            run = par.add_run(line)
            run.font.name = "Consolas"
            par.paragraph_format.left_indent = Pt(24)
            continue

        if not stripped:
            active_table = None
            continue

        # 表格
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                continue  # 表头分隔行
            if active_table is None:
                active_table = doc.add_table(rows=0, cols=len(cells))
            row = active_table.add_row()
            for i, c in enumerate(cells):
                if i >= len(row.cells):
                    break
                cell = row.cells[i]
                cell.text = ""
                _add_runs(cell.paragraphs[0], c)
            continue
        active_table = None

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            par = doc.add_heading(level=min(len(m.group(1)), 4))
            _add_runs(par, m.group(2))
            continue

        # 分隔线
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            par = doc.add_paragraph()
            _add_hr(par)
            continue

        # 有序列表
        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            par = doc.add_paragraph(style="List Number")
            _add_runs(par, m.group(1))
            continue

        # 无序列表
        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            par = doc.add_paragraph(style="List Bullet")
            _add_runs(par, m.group(1))
            continue

        # 普通段落
        par = doc.add_paragraph()
        _add_runs(par, stripped)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def report_markdown_to_pdf(markdown: str) -> bytes:
    """将 Markdown 报告内容转换为 PDF 二进制。

    使用 reportlab 渲染中文（内建 STSong-Light CID 字体，无需外部字体文件，
    离线安全）。支持标题、段落、无序/有序列表、| 表格、分隔线；
    内联 SVG 图表（data URI）无法嵌入 PDF，故按行跳过（明细表已保留在正文）。
    """
    import re
    from io import BytesIO
    from xml.sax.saxutils import escape as _xml_escape_pdf

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    _FONT = "STSong-Light"
    with contextlib.suppress(Exception):  # 字体已注册或字体表异常时不影响流程
        pdfmetrics.registerFont(UnicodeCIDFont(_FONT))

    _base = getSampleStyleSheet()
    styles = {}
    for name in ("Title", "Heading1", "Heading2", "Heading3", "Heading4", "Normal"):
        st = _base[name].clone(name=name)
        st.fontName = _FONT
        st.fontSize = 15 if name == "Title" else (13 if name == "Heading1" else 11)
        st.leading = st.fontSize * 1.5
        st.spaceAfter = 6
        if name == "Title":
            st.alignment = TA_CENTER
        styles[name] = st
    styles["Bullet"] = _base["Normal"].clone(name="Bullet")
    styles["Bullet"].fontName = _FONT
    styles["Bullet"].fontSize = 10.5
    styles["Bullet"].leading = 15
    styles["Bullet"].leftIndent = 12
    styles["Bullet"].bulletIndent = 0

    def _para(text, style, as_table_cell=False):
        # 转义 XML，再还原 Markdown 加粗/行内代码为 reportlab 标签
        safe = _xml_escape_pdf(text)
        safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
        safe = re.sub(r"`([^`]+)`", r"\1", safe)
        return Paragraph(safe, style)

    story = []
    in_code = False
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        # 内联图表（data URI）行跳过；代码块行不逐段渲染
        if in_code:
            if stripped.startswith("```"):
                in_code = False
            continue
        if stripped.startswith("```"):
            in_code = True
            continue
        if re.match(r"^!\[.*?\]\(data:image/svg\+xml;base64,", stripped):
            continue
        if not stripped:
            story.append(Spacer(1, 4))
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            story.append(Paragraph(_xml_escape_pdf(m.group(2)), styles[f"Heading{level}"]))
            continue

        # 表格
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                continue  # 分隔行
            table = Table([[Paragraph(c, styles["Normal"]) for c in cells]], hAlign="LEFT")
            table.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), _FONT),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f4f7")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c8cdd4")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(table)
            story.append(Spacer(1, 4))
            continue

        # 分隔线
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#999999")))
            continue

        # 有序 / 无序列表
        m = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if m:
            story.append(Paragraph(f"{m.group(1)}.&nbsp;{_xml_escape_pdf(m.group(2))}", styles["Normal"]))
            continue
        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            story.append(Paragraph(_xml_escape_pdf(m.group(1)), styles["Bullet"], bulletText="•"))
            continue

        # 普通段落
        story.append(_para(stripped, styles["Normal"]))

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=18 * mm, bottomMargin=18 * mm,
                            title="报告")
    doc.build(story)
    return buf.getvalue()

