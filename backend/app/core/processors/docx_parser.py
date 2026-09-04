"""DOCX 解析器：python-docx 取段落 + 表格（表格渲染为 GFM Markdown）。"""
import io

from .base import BaseParser, grid_to_markdown, register_parser


@register_parser(".docx")
class DocxParser(BaseParser):
    supported_ext = ".docx"

    def extract(self, file_bytes: bytes) -> str:
        import docx

        document = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        for tbl_idx, table in enumerate(document.tables, 1):
            rows = []
            for row in table.rows:
                cells = []
                prev = None
                for cell in row.cells:
                    # python-docx 合并单元格在同一行会重复返回同一对象，按对象去重
                    if cell is prev:
                        continue
                    prev = cell
                    cells.append(cell.text)
                rows.append(cells)
            md = grid_to_markdown(rows, f"### 表格 {tbl_idx}")
            if md:
                parts.append(md)
        return "\n".join(parts)