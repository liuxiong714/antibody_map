"""多格式文献文本提取分发器。

按文件扩展名分发到对应解析器，统一返回纯文本：
- .pdf  → PyMuPDF（复用 pdf_parser，含 OCR 兜底）
- .caj  → caj2pdf 转 PDF → PyMuPDF（尽力而为，失败抛 RuntimeError）
- .epub → ebooklib 读 XHTML → bs4 取文本
- .docx → python-docx 取段落
- .txt  → utf-8 解码（回退 gbk/gb18030）
- .html/.htm → bs4 去标签取文本

MIME 映射表供上传校验、文件服务、MinIO 上传共用。
"""
import io
import logging
import shutil
import subprocess
import tempfile
from pathlib import Path

from app.core.pdf_parser import extract_text as pdf_extract_text

logger = logging.getLogger("uvicorn")

# 扩展名 → MIME 映射（上传白名单即该表的 keys）
MIME_MAP = {
    ".pdf": "application/pdf",
    ".caj": "application/octet-stream",
    ".epub": "application/epub+zip",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".html": "text/html",
    ".htm": "text/html",
}

ALLOWED_EXTS = set(MIME_MAP.keys())

# 标题清理时需识别的已知后缀（顺序无关，仅取首个匹配）
_KNOWN_EXTS = (".pdf", ".caj", ".epub", ".docx", ".txt", ".html", ".htm")


def get_mime_type(file_ext: str) -> str:
    """按扩展名返回 MIME 类型，未知返回 application/octet-stream。"""
    ext = (file_ext or "").lower()
    if ext and not ext.startswith("."):
        ext = "." + ext
    return MIME_MAP.get(ext, "application/octet-stream")


def normalize_ext(file_ext: str) -> str:
    """规范化扩展名为带点小写形式，默认 .pdf。"""
    ext = (file_ext or ".pdf").lower()
    if not ext.startswith("."):
        ext = "." + ext
    return ext


def extract_text(file_bytes: bytes, file_ext: str = ".pdf") -> str:
    """按扩展名分发到对应解析器，提取纯文本。

    CAJ 转换失败时抛 RuntimeError（由上层捕获并标记 failed），
    其他格式解析异常返回空串（上层据此判断"文本为空"）。
    """
    ext = normalize_ext(file_ext)

    try:
        if ext == ".pdf":
            return pdf_extract_text(file_bytes)
        if ext == ".caj":
            pdf_bytes = _caj_to_pdf_bytes(file_bytes)
            return pdf_extract_text(pdf_bytes)
        if ext == ".epub":
            return _extract_epub(file_bytes)
        if ext == ".docx":
            return _extract_docx(file_bytes)
        if ext == ".txt":
            return _extract_txt(file_bytes)
        if ext in (".html", ".htm"):
            return _extract_html(file_bytes)
        raise ValueError(f"不支持的文件格式: {ext}")
    except RuntimeError:
        # CAJ 转换失败等显式错误，向上传播以便日志明确
        raise
    except Exception as e:
        logger.error(f"文件解析失败 ({ext}): {e}")
        return ""


def _extract_epub(file_bytes: bytes) -> str:
    """从 EPUB 提取文本：ebooklib 读 XHTML → bs4 取文本。"""
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    # 抑制 ebooklib 的.epub.epubread 非致命警告噪音
    logging.getLogger("ebooklib").setLevel(logging.ERROR)

    book = epub.read_epub(io.BytesIO(file_bytes))
    parts = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_body_content() or b"", "html.parser")
        text = soup.get_text(separator="\n")
        if text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    """从 DOCX 提取段落文本。"""
    import docx  # python-docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # 表格单元格文本
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _extract_txt(file_bytes: bytes) -> str:
    """纯文本解码：utf-8 → gbk → gb18030 → ignore。"""
    for enc in ("utf-8", "gbk", "gb18030"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _extract_html(file_bytes: bytes) -> str:
    """HTML 去标签取文本。"""
    from bs4 import BeautifulSoup

    html = None
    for enc in ("utf-8", "gbk"):
        try:
            html = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if html is None:
        html = file_bytes.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator="\n")


def _caj_to_pdf_bytes(caj_bytes: bytes) -> bytes:
    """调用 caj2pdf 将 CAJ 转 PDF，返回 PDF 字节。

    依赖系统已安装 caj2pdf 命令与 mutool（mupdf-tools）。
    任一缺失或转换失败时抛 RuntimeError，由上层标记提取 failed。
    """
    if shutil.which("caj2pdf") is None:
        raise RuntimeError(
            "未安装 caj2pdf 命令，无法转换 CAJ（需 pip install caj2pdf 并安装 mutool）"
        )
    if shutil.which("mutool") is None:
        logger.warning("未检测到 mutool，CAJ 转换可能失败（caj2pdf 依赖 mutool）")

    with tempfile.TemporaryDirectory() as tmpdir:
        caj_path = Path(tmpdir) / "input.caj"
        pdf_path = Path(tmpdir) / "output.pdf"
        caj_path.write_bytes(caj_bytes)
        try:
            result = subprocess.run(
                ["caj2pdf", "convert", str(caj_path), "-o", str(pdf_path)],
                capture_output=True,
                text=True,
                timeout=120,
            )
        except FileNotFoundError as e:
            raise RuntimeError(f"caj2pdf 命令调用失败: {e}")
        except subprocess.TimeoutExpired:
            raise RuntimeError("CAJ 转换超时（120s）")

        if not pdf_path.exists() or pdf_path.stat().st_size == 0:
            stderr = (result.stderr or "").strip()
            tail = f": {stderr[:300]}" if stderr else ""
            raise RuntimeError(
                f"CAJ 转换失败，可能为不支持的 CAJ 变体或加密文件{tail}"
            )
        return pdf_path.read_bytes()
